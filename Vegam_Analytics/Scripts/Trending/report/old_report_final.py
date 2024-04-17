import re
import fitz  # PyMuPDF
import os
# os.chdir(r'D:\vegam\onedrive_vegam\OneDrive - Aureole Technologies Private Limited\analytic_group\Trending\Report\script\ps\Reports')
os.chdir(r'D:\analytics_group\Reports')
import shutil
existing_file_path = "report_making.docx"
from docx.oxml import parse_xml
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Pt, RGBColor
# from docx.shared import Pt
import pandas as pd
from docx.shared import Inches, Cm
from docx.oxml import OxmlElement,ns
from docx.oxml.ns import qn
import json
import numpy as np
from scipy.stats import linregress
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('TkAgg')  # Use the TkAgg backend (or another suitable one)
import time
from datetime import datetime, timedelta
from docx.enum.style import WD_STYLE_TYPE



#first api call
common_import = {"site_id": "001",
                 "start_date": "01 November 2023",
                 "end_date": "14 November 2023"}  # will be imported -- json 1 [ common_import.json ]
import requests
import json


def api_call_for_metadata(api_url,
                          json_data=None):  # json_data is None by default here as for us it is only get call even if api method is post
    try:
        ''' #later activate-- do not delete
        if json_data is None:
            #GET request
            response = requests.get(api_url)
        else:
            #POST
            response = requests.post(api_url, json=json_data)

        response.status_code = 200 #for now, it will be removed later
        if response.status_code == 200:
            result_json = response.json()
            return result_json
        else:
            print(f"Failed to make API call. Status code: {response.status_code}")
            print(response.text)
            return None

        '''
        # Read sample payload
        with open("DeployedSensorAPI.json", "r") as json_file:
            sample_payload = json.load(json_file)
        return sample_payload


    except Exception as e:
        print(f"An error occurred in api_call_for_metadata method: {e}")
        return None


site_id = common_import["site_id"]
api_url_get = f"http://192.168.1.129:8082/DeployedSensorApi?site_id={site_id}"
result_get = api_call_for_metadata(api_url_get)  # later activate

if result_get is not None:
    print("GET API call successful. Received JSON:")
    print(result_get) #metadata json
else:
    print("GET API call failed.")

#got metadata json in result_get variable





def duplicate_and_add_title(original_path, duplicate_path, title_image_path):
    # Duplicate the original DOCX file
    shutil.copyfile(original_path, duplicate_path)

    # Open the duplicated file
    doc = Document(duplicate_path)

    # Add the "pm_ref_1" image as the title
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run()
    title_run.add_picture(title_image_path, width=Inches(3))#,height=Inches(2))  # Adjust the width as needed
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #
    # doc.add_page_break()
    doc.save(duplicate_path)

original_file_path = "template.docx"
duplicate_file_path = "report_making.docx"
title_image_path = "vegam_logo.jpg"

duplicate_and_add_title(original_file_path, duplicate_file_path, title_image_path)

# do not change anything here till now...

def add_title_to_existing_doc(existing_path, title):
    doc = Document(existing_path)
    title_paragraph = doc.add_paragraph()
    title_run = title_paragraph.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(26)  # Set font size, adjust as needed

    # Set the alignment of the title paragraph to center
    title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


    doc.add_page_break()
    doc.save(existing_path)

    #pm
    heading = doc.add_heading('\n\nCONTENTS', level=1)  # pm
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 105, 148)#(0, 0, 255)  # Set font color to black
        run.font.size = Pt(22)  # pm
        run.font.name = "Calibri (Body)"  # pm
    doc.add_page_break()
    doc.save(existing_path)
    #pm




customer_name =result_get["SiteName"]

def format_date(date_str):
    date_obj = datetime.strptime(date_str, "%d %B %Y")
    return date_obj.strftime("%d %b %Y")

formatted_start_date = format_date(common_import["start_date"])
formatted_end_date = format_date(common_import["end_date"])

duration = f"{formatted_start_date} – {formatted_end_date}"
# "SEP – OCT 2023"

# title_text = '\n\n\n\n\n       MBS KINGMAN \n     Condition Monitoring Report\n       SEP – OCT 2023'  #need to add logic for get data from excel
title_text = f'\n\n\n\n\n       {customer_name} \n     Condition Monitoring Report\n       {duration}'

add_title_to_existing_doc(existing_file_path, title_text)

# doc output_report_final.docx got created with fisrt page and 2nd content page but empty



#supriya code task 8
doc = Document(existing_file_path)
def add_summary(doc):
    # Add a heading without side margins
    summary_heading = doc.add_heading('Summary:', level=1)
    for run in summary_heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
        run.font.size = Pt(18)
        run.font.name = 'Calibri'  # Set font to Calibri
        run.bold = True  # Set to bold

    # Add a table without side margins
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"

    # Calculate the width of the cell based on the content
    content_width = Pt(900)  # Adjust the initial width as needed

    # content_text = (
    #     "Extruder Drive Gear Box 1301, Extruder Drive Gear Box 1201 hi i am purbasha"
    # )
    content_text = (
        "---------- This is manually written text ------------ \nExtruder Drive Gear Box 1301, Extruder Drive Gear Box 1201, "
        "and Gear pump Drive Gear Box 1202 machine vibration levels are observed within"
        " a good range, even though momentary spikes are found due to the process. "
        "Godet 1 Gear Box 1401 machine noted with high vibration range & equipment’s"
        " parameters are required for further fault analysis. \n -----------end of manually written text ------------"
        "\n [content of the text which will be changed dynamically based on VRMS data --> task yet to be start/ added #pm]"
    )
    cell_paragraph = table.cell(0, 0).paragraphs[0]
    cell_paragraph.text = content_text

    # Set the calculated width to the cell
    table.columns[0].width = content_width

    # Merge cells to create the rectangle shape
    table.cell(0, 0).merge(table.cell(0, 0))

    # Add a text box-like effect using a paragraph in the cell
    cell_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell_paragraph.style ="Normal" #"ExistingStyle" #"BodyText" #pm

    # doc.styles.add_style('BodyText', WD_PARAGRAPH_ALIGNMENT.LEFT, Pt(12))
    #
    # cell_paragraph.style = "BodyText"

    # Adjust the font size based on the width
    font_size = Pt(14)
    cell_paragraph.runs[0].font.size = font_size
    while cell_paragraph.runs[0].font.size > Pt(8):
        if cell_paragraph.runs[0].font.size * len(content_text) > content_width:
            font_size -= Pt(1)
            cell_paragraph.runs[0].font.size = font_size
        else:
            break

    cell_paragraph.text = content_text
    cell_paragraph.runs[0].font.bold = False  # Set to not bold

    # Calculate the height of the cell based on content
    content_height = len(content_text.split("\n")) * font_size * 1.2  # Adjust the factor as needed
    table.cell(0, 0).height = content_height
    border_xml = """
        <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
            <w:top w:val="single" w:color="89CFF0" w:sz="40" w:space="0" w:shadow="false" w:frame="true"/>
            <w:bottom w:val="single" w:color="89CFF0" w:sz="40" w:space="0" w:shadow="false" w:frame="true"/>
            <w:left w:val="single" w:color="89CFF0" w:sz="40" w:space="0" w:shadow="false" w:frame="true"/>
            <w:right w:val="single" w:color="89CFF0" w:sz="40" w:space="0" w:shadow="false" w:frame="true"/>
            <w:shd w:fill="{shading_color}" w:val="clear" w:color="auto"/>
        </w:tcBorders>
    """
    border_element = parse_xml(border_xml)
    cell = table.cell(0, 0)

    cell._element.tcPr.append(border_element)

#Overall Equipment Overview Table
# Function to extract 'EquipmentID' and 'SensorMacID' and combine them into 'EquipmentID_Sensor_ID'
def extract_equipment_sensor_ids(data):
    equipment_sensor_list = []

    def extract(data):
        if isinstance(data, dict):
            if 'EquipmentID' in data and 'AttachedSensors' in data:
                for sensor in data['AttachedSensors']:
                    equipment_id = data['EquipmentID']
                    sensor_mac_id = sensor['SensorMacID']
                    equipment_sensor_list.append(f"{equipment_id}_{sensor_mac_id}")
            for value in data.values():
                extract(value)
        elif isinstance(data, list):
            for item in data:
                extract(item)

    extract(data)
    return equipment_sensor_list

# # Get a list of combined 'EquipmentID' and 'SensorMacID'
# combined_equipment_sensor_ids = extract_equipment_sensor_ids(result_get)

def equipment_df(combined_equipment_sensor_ids):
    # Create a DataFrame with 'EquipmentID_Sensor_ID' column
    df = pd.DataFrame({'EquipmentID_Sensor_ID': combined_equipment_sensor_ids})
    df = df[['EquipmentID_Sensor_ID']]  # Select only the 'EquipmentID_Sensor_ID' column
    return df

def start_end_duration(common_config_file):
    

    start_time_str = common_config_file['start_time']
    end_time_str = common_config_file['end_time']
    
    # Convert string timestamps to datetime objects
    start_time = datetime.strptime(start_time_str, '%d/%m/%y %I:%M %p')
    end_time = datetime.strptime(end_time_str, '%d/%m/%y %I:%M %p')
    
    # Extract the days between start and end time
    duration = end_time - start_time
    
    # Format start and end dates to show day, month, and year
    formatted_start_date = start_time.strftime('%d %b %Y')
    formatted_end_date = end_time.strftime('%d %b %Y')
    
    print(f"Start Date: {formatted_start_date}")
    print(f"End Date: {formatted_end_date}")

    # Extract the total number of days between start and end time
    total_days = duration.days
    
    print(f"Total Duration: {total_days} days")
    
    return start_time, end_time, formatted_start_date, formatted_end_date, total_days


def create_table_with_header(doc, days_list, num_days, equipment_sensor_ids):
    num_days = len(days_list)
    table = doc.add_table(rows=len(equipment_sensor_ids) + 2, cols=num_days + 1)
    table.style = 'Table Grid'

    equipment_ids = equipment_sensor_ids['EquipmentID_Sensor_ID'].tolist()

    # Set the width of the 'Equipment' column
    equipment_column = table.columns[0]
    equipment_column.width = Inches(2.5)  # Change the width as needed


    # Fill the second row headers for days
    for day, date in enumerate(days_list):
        table.cell(0, day + 1).text = date

    # Fill the first column headers for equipment
    for equip, equipment_id in enumerate(equipment_ids):
        table.cell(equip + 1, 0).text = f"{equipment_id}"

def add_column_names_to_subsequent_pages(doc, num_days):
    # Loop through the sections after the first one (as the first section's header is set separately)
    for section in doc.sections[1:]:
        # header = section.header
        table = doc.tables[0]  # Assuming the table is the first table in the document
        
        # Add column names to the header of subsequent pages
        for day in range(num_days):
            header_cell = table.cell(0, day + 1)
            new_paragraph = header_cell.paragraphs[0].add_run(f"{day + 1}")
            new_paragraph.font.bold = False  # Bold for column names in the header
            new_paragraph.font.size = Pt(5)

def add_header_outside_table(doc, header_text):
    # Add a paragraph before the table to serve as the header
    header_paragraph = doc.add_paragraph()
    header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    header_run = header_paragraph.add_run(header_text)
    header_run.bold = True
    header_run.font.size = Inches(0.2)  # Adjust font size as needed


# Functions for Equipment Overview Status Table
def add_equipment_overview_status_table(doc, num_rows):
    main_table = doc.add_table(rows=num_rows + 1, cols=4)
    main_table.style = 'Table Grid'

    main_column_names = main_table.rows[0].cells
    main_column_names[0].text = 'Sl. No'
    main_column_names[1].text = 'EQUIPMENT ID'
    main_column_names[2].text = 'STATUS'
    main_column_names[3].text = 'COMMENTS'

    main_table.cell(0, 0).width = Inches(1.0)
    main_table.cell(0, 1).width = Inches(5.0)
    main_table.cell(0, 2).width = Inches(5.0)
    main_table.cell(0, 3).width = Inches(5.0)
    for cell in main_table.rows[0].cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(12)



    # Fill the "Sl. No" column dynamically
    for row_num in range(1, num_rows + 1):
        main_table.cell(row_num, 0).text = str(row_num)

        for paragraph in main_table.cell(row_num, 0).paragraphs:
            # Remove bold formatting from the entire row
            for run in paragraph.runs:
                run.bold = False
                # Maintain font size at 11 for all runs
                run.font.size = Pt(11)
    return main_table # pm



def check_and_add_equipment_ids(main_table, equipment_ids_df):
    # Convert target column name to lowercase
    target_column_name = "equipment id"

    # Check if the target column name exists in the Excel file (case insensitive)
    equipment_id_column_name = next(
        (col_name for col_name in equipment_ids_df.columns if col_name.lower() == target_column_name),
        None
    )

    sensor_id_column_name = next(
        (col_name for col_name in equipment_ids_df.columns if col_name.lower() == "sensor id"),
        None
    )

    # Check if the columns exist in the Excel file (case insensitive)
    if not equipment_id_column_name or not sensor_id_column_name:
        raise ValueError("Could not find required columns in Excel file")

    # If the columns exist, add the combined data to the "EQUIPMENT ID_SENSOR ID" column in main_table
    for row_num, row in equipment_ids_df.iterrows():
        equipment_id = row[equipment_id_column_name]
        sensor_id = row[sensor_id_column_name]
        combined_equipment_id = f"{equipment_id}_{sensor_id}"
        main_table.cell(row_num + 1, 1).text = str(combined_equipment_id)

        for paragraph in main_table.cell(row_num + 1, 1).paragraphs:
            # Remove bold formatting from the entire row
            for run in paragraph.runs:
                run.bold = False
                # Maintain font size at 11 for all runs
                run.font.size = Pt(11)

    # # If the column exists, add the data to the "EQUIPMENT ID" column in main_table
    # if equipment_id_column_name:
    #     for row_num, equipment_id in enumerate(equipment_ids_df[equipment_id_column_name]):
    #         main_table.cell(row_num + 1, 1).text = str(equipment_id)
    #
    #         for paragraph in main_table.cell(row_num + 1, 1).paragraphs:
    #             # Remove bold formatting from the entire row
    #             for run in paragraph.runs:
    #                 run.bold = False
    #                 # Maintain font size at 11 for all runs
    #                 run.font.size = Pt(11)


def add_key_table_heading(doc):
    key_heading_paragraph = doc.add_paragraph('Key:')
    key_heading_paragraph.style.font.bold = True  # Make the "Key" text bold
    key_heading_paragraph.style.font.size = Pt(18)#pm ch
    key_heading_paragraph.style.font.name = "Calibri (Body)"  # pm
    # run.font.size = Pt(18)
    # run.font.name = "Calibri (Body)"  # pm


def add_key_table(doc):
    # Add a table for the key information

    key_table = doc.add_table(rows=3, cols=4)
    key_table.style = 'Table Grid'


    key_table.cell(0, 0).width = Inches(1.0)
    key_table.cell(0, 1).width = Inches(7.0)
    key_table.cell(0, 2).width = Inches(1.0)
    key_table.cell(0, 3).width = Inches(7.0)

    # Fill table with text
    key_table.cell(0, 0).text = ''
    key_table.cell(0, 1).text = 'GOOD'
    key_table.cell(0, 2).text = ''
    key_table.cell(0, 3).text = 'CRITICAL'

    key_table.cell(1, 0).text = ''
    key_table.cell(1, 1).text = 'Fluctuating / Unknown'
    key_table.cell(1, 2).text = ''
    key_table.cell(1, 3).text = 'Machine NOT RUNNING'

    key_table.cell(2, 0).text = ''
    key_table.cell(2, 1).text = 'WARNING'
    key_table.cell(2, 2).text = ''
    key_table.cell(2, 3).text = 'Data is not available'

    # Adjust font size
    for row in key_table.rows:
        for cell in row.cells:
            if not cell.paragraphs:
                cell.add_paragraph('')
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(12)
                cell.paragraphs[0].runs[0].bold = True
                if len(cell.paragraphs[0].runs) > 1:
                    cell.paragraphs[0].runs[1].font.size = Pt(12)

    # Apply shading to specific rows and columns in the key table
    apply_shading(key_table.cell(0, 0), "32CD32")
    apply_shading(key_table.cell(0, 2), "FF0000")
    apply_shading(key_table.cell(1, 0), "FFFF00")
    apply_shading(key_table.cell(1, 2), "000000")
    apply_shading(key_table.cell(2, 0), "FFA500")
    apply_shading(key_table.cell(2, 2), "FFFFFF")
    return key_table

# def apply_shading(cell, color):
#     cell_xl_element = cell._element
#     tab_cell_prop = cell_xl_element.get_or_add_tcPr()
#
#     # Create shading object
#     shade_obj = OxmlElement('w:shd')
#     # Set the shading color
#     shade_obj.set(qn('w:fill'), color)
#     # Append the shading properties to the cell properties
#     tab_cell_prop.append(shade_obj)

cell_xl_element = None
def apply_shading(cell, color):
    global cell_xl_element
    cell_xl_element = cell._element
    tab_cell_prop = cell_xl_element.get_or_add_tcPr()

    # Create shading object
    shade_obj = OxmlElement('w:shd')
    # Set the shading color
    shade_obj.set(qn('w:fill'), color)
    # Append the shading properties to the cell properties
    tab_cell_prop.append(shade_obj)

add_summary(doc)
doc.save(existing_file_path)
# exit()

#Calling Over All Overview Status Table
doc = Document(existing_file_path)

time_duration = start_end_duration(common_import)

start_time = time_duration[0]
end_time = time_duration[1]
formatted_start_date = time_duration[2]
formatted_end_date = time_duration[3]
total_days = time_duration[4]

# Initialize a list to store the days
days_list = []

# Calculate the days between start and end dates
current_date = start_time
while current_date < end_time + timedelta(days=1):
    days_list.append(current_date.strftime('%d'))
    current_date += timedelta(days=1)


# Add a header above the table
header_text = f"Overall Plant Status ({formatted_start_date} to {formatted_end_date}):"
add_header_outside_table(doc, header_text)

# The duration in days
duration_days = total_days  # Change this value to set the duration in days (columns)

# Calculate the number of days columns dynamically
num_days = duration_days

# Get a list of combined 'EquipmentID' and 'SensorMacID'
combined_equipment_sensor_ids = extract_equipment_sensor_ids(result_get)

equip_details_df =  equipment_df(combined_equipment_sensor_ids)
# # Create a table with 'num_equipment' rows and 'num_days' + 1 columns
# create_table_with_header(doc, num_days, df)
create_table_with_header(doc, days_list, num_days, equip_details_df)

# Add column names to subsequent pages
add_column_names_to_subsequent_pages(doc, num_days)

# Save the document
doc.save(existing_file_path)



heading = doc.add_heading('2. Equipment Overview Status:', level=1)
for run in heading.runs:
    run.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black
    run.font.size = Pt(18) #pm
    run.font.name = "Calibri (Body)" #pm


excel_file_path = 'Inputs_for_report.xlsx'  # Replace with your actual Excel file path
equipment_ids_df = pd.read_excel(excel_file_path)

# Dynamically determine the number of rows based on the Excel file
num_rows = len(equipment_ids_df)

# Add main table with the determined number of rows
print(doc, num_rows)
main_table = add_equipment_overview_status_table(doc, num_rows)
print(main_table)

# Check and add EQUIPMENT IDs to the main_table
check_and_add_equipment_ids(main_table, equipment_ids_df)
#new


doc.add_paragraph("\n")

# Add a key section
add_key_table_heading(doc)
add_key_table(doc)

# Save the modified document
# doc.save("result_doc.docx")
# doc.add_page_break() #pm # this was extra
doc.save(existing_file_path)


target_column_name = "equipment id"

# Check if the target column name exists in the Excel file (case insensitive)
equipment_id_column_name = next(
    (col_name for col_name in equipment_ids_df.columns if col_name.lower() == target_column_name),
    None
)
#pm modify


# Create an empty document
document = Document(existing_file_path) #pm

sections = document.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(1.91)
    section.right_margin = Cm(1.91)

# excel_file_name = "empty_document1"


document.save(existing_file_path)#save(r'D:\\analytics_group\\Reports\\' + '\\' + excel_file_name + '.docx')


##### Read Excel ####


input_excel_file = r'Inputs_for_report.xlsx'

input_df = pd.read_excel(input_excel_file)


#####################################################
# Add Heading
#####################################################
def add_heading_to_document(doc, text, level, color, font_name, font_size):
    heading = doc.add_paragraph(text)
    heading.style = f"Heading {level}"
    run = heading.runs[0]
    run.bold = True
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(*color)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Adjust alignment as needed

def add_table_with_empty_column(doc, data, image_path):
    table = doc.add_table(rows=len(data), cols=3, style='Table Grid')

    # Fill the first two columns with data
    for i, (text1, text2) in enumerate(data):
        row = table.rows[i]
        row.cells[0].text = text1
        row.cells[1].text = text2

    # Merge cells in the third column for rows 1 to 3
    for i in range(2):
        table.cell(i, 2).merge(table.cell(i + 1, 2))


    for table in doc.tables:
        for row in table.rows:
            row.height = Cm(1.0)
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    for p in cell.paragraphs:
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

                    for run in paragraph.runs:
                        run.font.name = "Calibri (Body)"
                        run.font.bold = True
                        run.font.size = Pt(12)

    # Insert image into the third column
    cell = table.cell(0, 2)  # Choose the position for the image insertion
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(2.5))  # Adjust image width as needed

    return table

# Function to load, preprocess data and create plots
def preprocess_data(directory_path, file_name):
    # if dataframes is None:
    dataframes = []  # Initialize the dictionary if it doesn't exist

    # # Calculate the relative path based on the current file being processed
    data = directory_path + '\\' + file_name +'.xlsx'
    df = pd.read_excel(data, sheet_name=None)  # Read all sheets into a dictionary
    for axis_name, sheet_name in [('X_Axis', 'Series_1'), ('Y_Axis', 'Series_2'), ('Z_Axis', 'Series_3')]:
        df_axis = df.get(sheet_name, pd.DataFrame()).copy()  # Create a copy

        # Select the columns you want to keep (e.g., the last two columns)
        columns_to_keep = ['DateTime', 'Value']  # Adjust column names as needed

        # Filter the DataFrame to include only the specified columns
        df_axis = df_axis[columns_to_keep]

        # Convert DateTime column to datetime format
        df_axis['DateTime'] = pd.to_datetime(df_axis['DateTime'], format='%d/%m/%Y %I:%M %p')
        # Store the DataFrame in the dictionary with its name
        dataframe_name = f'{file_name}_{axis_name}'.replace('.xls', '')
        dataframes.append((file_name, directory_path, axis_name, df_axis))

        print(f'Generated dataframe name: {dataframe_name}')

    return dataframes

def creating_directory(folder, list_of_sub_directories):
    for name in list_of_sub_directories:
        sub_diretories = folder + '\\' + name
        os.makedirs(sub_diretories, exist_ok=True)

def moving_average_trend(df, window_sizes):
    result_df = df.copy()

    for window_size in window_sizes:

        # Calculate the moving average
        result_df[f'MovingAverage_{window_size}'] = df['Value'].rolling(window=window_size).mean()
        result_df.loc[:window_size - 1, f'MovingAverage_{window_size}'] = df.loc[:window_size - 1, 'Value']  # pm

    return result_df

def trend_line_method(df):
    x = df['DateTime'].astype(np.int64) // 10 ** 9  # Convert to Unix timestamp (seconds)
    slope, intercept, _, _, _ = linregress(x, df['Value'])
    trendline = intercept + slope * x

    return trendline, slope

def moving_avg_trendline_threshold_plot(original_df, dataframe,
                             dataframe_name, df_name, config,
                             data_duration_days):

    trending_methods = config['trending_methods']
    for method in trending_methods:

        if method['enabled'] and method['name'] == 'Moving Average and Trendline':

            # Now, apply the trendline method to the original DataFrame
            trendline, slope = trend_line_method(original_df)
            print('Trend Line Data')
            window_sizes = method['burst_per_day']
            colors = ['magenta']

            threshold_value = config['Threshold_Value']
            for T_value in threshold_value:
                # if T_value['Warning'] == 'Warning':
                warning = T_value['Warning']
                # if T_value['Critical'] == 'Critical':
                critical = T_value['Critical']

            moving_average_dfs = {}

            plt.figure(figsize=(12, 6))

            # Calculate moving averages for each window size
            for window_size, color in zip(window_sizes, colors):
                result_df = moving_average_trend(dataframe.copy(), [window_size])
                result_df.fillna(0)
                moving_average_dfs[window_size] = result_df  # Store the result_df in the moving_average_dfs dictionary

                y_column_name = f'MovingAverage_{window_size}'

                value = max(result_df[y_column_name][window_size:])
                if value >= warning or warning <= value <= critical:
                    plt.axhline(y=warning, color='orange', linestyle='--', label='Warning')
                    plt.axhline(y=critical, color='red', linestyle='--', label='Critical')
                elif value <= warning:
                    plt.axhline(y=warning, color='orange', linestyle='--', label='Warning')

                # Plot the data
                plt.plot(result_df['DateTime'][window_size:], result_df[y_column_name][window_size:], color=color, label='Moving Average')
            original_df['Trendline'] = trendline


            plt.plot(original_df['DateTime'], original_df['Trendline'], color = 'lawngreen', label='Trendline')
            plt.xlabel('Date and Time')
            plt.ylabel('VelocityRMS (mm/sec)')
            plt.legend(loc='upper right')
            # title_name = dataframe_name[:-5]+f'_{df_name}'

            # Save the trendline chart as an image in the specified output format
            trendline_output_file = os.path.join(config['Chart_path'], f'{dataframe_name}_{df_name}.{config["output_format"]}')
            plt.savefig(trendline_output_file, format=config['output_format'])

            # Close the plot to release resources (optional)
            plt.close()

def moving_avg_trendline_df(original_df, dataframe, config):

    trending_methods = config['trending_methods']
    for method in trending_methods:

        if method['enabled'] and method['name'] == 'Moving Average and Trendline':

            # Now, apply the trendline method to the original DataFrame
            trendline, slope = trend_line_method(original_df)
            print('Trend Line Data')
            window_sizes = method['burst_per_day']


            moving_average_dfs = {}


            # Calculate moving averages for each window size
            for window_size in window_sizes:
                result_df = moving_average_trend(dataframe.copy(), [int(window_size)])
                result_df.fillna(0)
                moving_average_dfs[window_size] = result_df  # Store the result_df in the moving_average_dfs dictionary

                Moving_Average = f'MovingAverage_{window_size}'

            return trendline, result_df[Moving_Average]

def moving_avg_trendline_dataframe(dataframes, config):  # to plots and get pics

    file_name, item_path, axis_name, dataframe = dataframes
    # print(f'Processing dataframe: {file_name}')

    # # Get the current time
    # time = datetime.now()

    # # Format the current time as per the specified format
    # current_time = str(time.strftime('%d/%m/%Y %I:%M %p'))

    dataframe = dataframe

    original_df = dataframe

    Trendline, Moving_Average = moving_avg_trendline_df(original_df, dataframe, config)

    return Trendline, Moving_Average

def plotting(dataframes, output_folder, config):  # to plots and get pics

    # for file_name, item_path, axis_name, dataframe in dataframes:
    file_name, item_path, axis_name, dataframe = dataframes
    print(f'Processing dataframe: {file_name}')

    # # Get the current time
    # current_time = datetime.now()

    # # Format the current time as per the specified format
    # current_time = str(current_time.strftime('%d/%m/%Y %I:%M %p'))


    # dataframe = dataframe[(dataframe['DateTime'] <= current_time)]
    df = dataframe
    start_date = pd.to_datetime(df['DateTime'].min())
    end_date = pd.to_datetime(df['DateTime'].max())
    data_duration_days = ((end_date - start_date).days) + 1

    # # original_df = dataframe['DateTime']
    # original_df = dataframe[(dataframe['DateTime'] <= current_time)]
    original_df = dataframe

    moving_avg_trendline_threshold_plot(original_df, dataframe,
                                 file_name, axis_name, config,
                                 data_duration_days)

def threshold_value(value):
    if value < 5:
        return "Good"
    elif 5 <= value <= 8:
        return "Warning"
    else:
        return "Critical"

def overall_condition(df):
    has_warning = False
    has_critical = False

    for column in df:
        max_value = df[column].max()  # Find the maximum value in each column
        result = threshold_value(max_value)

        if result == "Warning":
            has_warning = True
        elif result == "Critical":
            has_critical = True
            break

    if has_critical:
        return "Critical"
    elif has_warning:
        return "Warning"
    else:
        return "Good"

def analyze_trend(data):
    first_value = data[0]
    last_value = data[len(data) - 1]

    thirty_percent_less = last_value * 0.3
    thirty_percent_more = last_value + (last_value * 0.3)
    fifty_percent_more = last_value + (last_value * 0.5)

    if first_value <= thirty_percent_less:
        return "Up trend"
    elif first_value >= thirty_percent_more:
        return "No uptrend"
    elif first_value >= fifty_percent_more:
        return "Down trend"
    else:
        return "No clear trend"

def analyze_condition(condition):
    if condition == "Good":
        return "Moving average trend appears to be normal and within ISO threshold limit during the process."
    elif condition == "Warning":
        return "Moving average trend appears to be warning."
    elif condition == "Critical":
        return "Moving average trend requires attention as it's beyond the ISO threshold limit during the process."
    # else:
    #     return "Condition not recognized. Please check input."

def check_overall_conditions(conditions):
    if 'Critical' in conditions:
        return 'Critical'
    elif 'Warning' in conditions:
        return 'Warning'
    else:
        return 'Good'

config_file_path = 'config (1).json'

# Load configuration from JSON file
config_file = config_file_path
with open(config_file, 'r') as f:
    config = json.load(f)


def generate_comments(list_1, list_2):
    warning_comments = []
    critical_comments = []

    for idx, status in enumerate(list_2):
        if status == 'Warning':
            warning_comments.append(list_1[idx])
        elif status == 'Critical':
            critical_comments.append(list_1[idx])

    if 'Warning' not in list_2 and 'Critical' not in list_2:
        return 'All three axis vibrations are observed within ISO threshold limit.'

    if len(warning_comments) == 1 and len(critical_comments) == 1:
        return f"warning range for {list_1[list_2.index('Warning')]} direction and critical range for {list_1[list_2.index('Critical')]} direction"

    if len(warning_comments) == 1 and len(critical_comments) == 2:
        return f"warning range for {list_1[list_2.index('Warning')]} direction and critical range for {critical_comments[1]} and {list_1[list_2.index('Critical')]} direction"

    if len(warning_comments) == 2 and len(critical_comments) == 1:
        return f"warning range for {warning_comments[1]} and {list_1[list_2.index('Warning')]} direction and critical range for {list_1[list_2.index('Critical')]} direction"

    if len(warning_comments) == 3 and len(critical_comments) == 0:
        return f"warning range for {warning_comments[1]}, {warning_comments[2]} and {list_1[list_2.index('Warning')]} direction"

    if len(warning_comments) == 0 and len(critical_comments) == 3:
        return f"critical range for {critical_comments[1]}, {critical_comments[2]} and {list_1[list_2.index('Critical')]} direction"

    if warning_comments:
        return f"Warning: Vibrations are observed in the warning range for the {', '.join(warning_comments)} direction as per ISO standards and require continuous monitoring."

    if critical_comments:
        return f"Critical: Vibrations are observed in the critical range for the {', '.join(critical_comments)} direction as per ISO standards and require immediate Service."

# Example usage
list_1 = ['Radial Horizontal', 'Axial', 'Radial Vertical']
list_2 = ['Good', 'Good', 'Good']

result = generate_comments(list_1, list_2)
print(result)

def remark_text(condition):
    if condition == 'Good':
        return ("*Overall condition of the machine/motor’s condition is good as the vibration data are observed within ISO threshold limit.\n"
        "* To enable detailed analysis and fault alerts, machine parameters related to rpm, bearing, and gear are required.")
    elif condition == 'Warning':
        return ("*Overall condition of the machine/motor’s condition is warning as the VRMS data are observed in the warning range as per the ISO threshold limits.\n"
        "* To enable detailed analysis and fault alerts, machine parameters related to rpm, bearing, and gear are required.")
    elif condition == 'Critical':
        return ("*Overall condition of the machine/motor’s condition is critical as the VRMS data are observed in the Critical range as per the ISO threshold limits..\n"
        "* To enable detailed analysis and fault alerts, machine parameters related to rpm, bearing, and gear are required.")
    else:
        return ("Condition not recognized. Please provide 'good', 'warning', or 'critical'.")



# remark_text = (
#     "*Overall condition of the machine/motor’s condition is good as the vibration data are observed within ISO threshold limit.\n"
#     "* To enable detailed analysis and fault alerts, machine parameters related to rpm, bearing, and gear are required."

# )

def X_Axis(input_df, axes, equipment_id, existing_file_path, config_file_path):
    filtered_data = input_df[input_df['Equipment ID'] == equipment_id][['SL No','Equipment ID','Duration', axes, 'Temp (in Degrees)']]
    slno = str(filtered_data['SL No'].item())
    Equipment_ID = equipment_id
    Duration = str(filtered_data['Duration'].item())
    Axis_name = str(filtered_data[axes].item())
    Temp = str(filtered_data['Temp (in Degrees)'].item())

    # Add Heading
    doc_path = existing_file_path

    # Load the existing document
    document = Document(doc_path)

    # Add a main heading with black color and specified font to the existing document
    main_heading_text = "3. Equipment Details:"
    main_heading_color = (0, 0, 0)  # RGB color values for black
    main_heading_font = "Calibri (Body)"  # Font name
    font_size = 18
    add_heading_to_document(document, main_heading_text, 1, main_heading_color,
                            main_heading_font, font_size)

    # Add a subheading with black color and specified font to the existing document
    sub_heading_text = "3."+slno+' '+Equipment_ID +':'
    sub_heading_color = (0, 0, 0)  # RGB color values for black
    sub_heading_font = "Calibri (Body)"  # Font name
    font_size = 16
    add_heading_to_document(document, sub_heading_text, 2, sub_heading_color,
                            sub_heading_font, font_size)


    # Save the updated document
    document.save(doc_path)

    # Get the 'data_Input_folder' from the config
    directory_path = config['data_Input_folder']

    # Create an output folder if it doesn't exist
    output_folder = config['Chart_path']
    # Call the preprocess_data function with the directory path
    dataframes = preprocess_data(directory_path, Equipment_ID)

    dataframes = dataframes[0]

    # Iterate through the dataframes and apply selected trending methods
    plotting(dataframes, output_folder, config)

    Trendline, Moving_Avg = moving_avg_trendline_dataframe(dataframes, config)

    TL_MA_df = pd.DataFrame({'Trend_Line': Trendline, 'Moving_Average': Moving_Avg})

    #Check the health of the equipment(Data:Moving average and TrendLine)
    overall_result = overall_condition(TL_MA_df)

    # Example data for the first two columns
    data = [
        ("Equipment Condition", overall_result),
        ("Duration", Duration),
        ("Highest Surface Temperature", Temp +" °C")
    ]
    # Load the existing document
    document = Document(doc_path)

    machine_image_path = config['machine_image_path']+'\\'+Equipment_ID+'.png'
    add_table_with_empty_column(document, data, machine_image_path)

    # Save the updated document
    document.save(doc_path)

    #Velocity RMS Data Summary Table

    # Add a subheading
    sub_heading_text = f"Velocity RMS Data Summary - {Axis_name}:"
    sub_heading_color = (0, 0, 0)  # RGB color values for black
    sub_heading_font = "Calibri (Body)"  # Font name
    font_size = 14
    add_heading_to_document(document, sub_heading_text, 2, sub_heading_color,
                            sub_heading_font, font_size)

    table = document.add_table(rows=2, cols=1, style='Table Grid')


    # Set the first row height to be three times larger than the second row
    for index, row in enumerate(table.rows):
        if index == 0:
            row.height = Cm(10)  # Set the first row height (in inches)
        else:
            row.height = Cm(4)  # Set the second row height (1/3 of the first row)



    # Access the cell at position [0, 0] (first row, first column)
    cell = table.cell(0, 0)

    # Add an image to the cell
    image_path = config['Chart_path']+'\\'+Equipment_ID+'_'+'X_Axis'+'.jpeg' # Replace with the path to your image file
    cell_paragraph = cell.paragraphs[0]
    run = cell_paragraph.add_run()
    run.add_picture(image_path, width=Cm(17.75))  # Adjust width as needed

    moving_average_observations = analyze_condition(overall_result)
    trend_observations = analyze_trend(Trendline)

    content = [
        {"text": "Observations:", "bold": True},  # Bold "Observations"
        {
            "text": moving_average_observations,
            "bullet": True,
            "spacing": True
        },  # Content with bullet point and spacing
        {
            "text": "",
            "bullet": False,
            "spacing": True
        },  # Content with bullet point and no spacing
        {
            "text": f"{trend_observations} is observed in the VRMS trend rate chart over the period",
            "bullet": True,
            "spacing": True
        }  # Content with bullet point and no spacing
    ]

    second_row = table.rows[1]  # Accessing the second row
    cell = second_row.cells[0]  # Accessing the first cell in the second row

    # Add the content to the cell as paragraphs
    for para_content in content:
        p = cell.add_paragraph()

        if isinstance(para_content, dict) and para_content.get("bold"):
            run = p.add_run(para_content["text"])
            run.bold = True  # Bold the text
            run.font.name = 'Calibri'  # Set font style to Calibri (Body)
            run.font.size = Pt(12)  # Set font size
        else:
            run = p.add_run(para_content["text"])
            run.font.name = 'Calibri'  # Set font style to Calibri (Body)
            run.font.size = Pt(12)  # Set font size

        # # style = document.styles.get('List Bullet')
        # if para_content.get("bullet"):
        #     p.style = 'List Bullet 2'  # Apply bullet point style

        p_format = p.paragraph_format
        if para_content.get("spacing"):
            p_format.space_before = Pt(5)  # 8 pt spacing before the paragraph

        # Add indentation for bullet points
        if para_content.get("bullet"):
            run = p.runs[0]
            run.text = " " * 2 + " " * 1 + run.text  # Adding 6 spaces before and 4 spaces after the bullet


    # Save the updated document
    document.save(doc_path)


def Y_Axis(input_df, axes, equipment_id, existing_file_path, config_file_path):
    filtered_data = input_df[input_df['Equipment ID'] == equipment_id][['SL No','Equipment ID','Duration', axes, 'Temp (in Degrees)']]
    Equipment_ID = equipment_id
    Duration = str(filtered_data['Duration'].item())
    Axis_name = str(filtered_data[axes].item())
    Temp = str(filtered_data['Temp (in Degrees)'].item())

    # Add Heading
    doc_path = existing_file_path

    # Load the existing document
    document = Document(doc_path)

    # Get the 'data_Input_folder' from the config
    directory_path = config['data_Input_folder']

    # Create an output folder if it doesn't exist
    output_folder = config['Chart_path']
    # Call the preprocess_data function with the directory path
    dataframes = preprocess_data(directory_path, Equipment_ID)

    dataframes = dataframes[1]

    # Iterate through the dataframes and apply selected trending methods
    plotting(dataframes, output_folder, config)

    Trendline, Moving_Avg = moving_avg_trendline_dataframe(dataframes, config)

    TL_MA_df = pd.DataFrame({'Trend_Line': Trendline, 'Moving_Average': Moving_Avg})

    #Check the health of the equipment(Data:Moving average and TrendLine)
    overall_result = overall_condition(TL_MA_df)

    # Save the updated document
    document.save(doc_path)

    #Velocity RMS Data Summary Table

    # Add a subheading
    sub_heading_text = f"Velocity RMS Data Summary - {Axis_name}:"
    sub_heading_color = (0, 0, 0)  # RGB color values for black
    sub_heading_font = "Calibri (Body)"  # Font name
    font_size = 14
    add_heading_to_document(document, sub_heading_text, 2, sub_heading_color,
                            sub_heading_font, font_size)

    table = document.add_table(rows=2, cols=1, style='Table Grid')


    # Set the first row height to be three times larger than the second row
    for index, row in enumerate(table.rows):
        if index == 0:
            row.height = Cm(10)  # Set the first row height (in inches)
        else:
            row.height = Cm(4)  # Set the second row height (1/3 of the first row)

    # Access the cell at position [0, 0] (first row, first column)
    cell = table.cell(0, 0)

    # Add an image to the cell
    image_path = config['Chart_path']+'\\'+Equipment_ID+'_'+'Y_Axis'+'.jpeg' # Replace with the path to your image file
    cell_paragraph = cell.paragraphs[0]
    run = cell_paragraph.add_run()
    run.add_picture(image_path, width=Cm(17.75))  # Adjust width as needed

    moving_average_observations = analyze_condition(overall_result)
    trend_observations = analyze_trend(Trendline)

    content = [
        {"text": "Observations:", "bold": True},  # Bold "Observations"
        {
            "text": moving_average_observations,
            "bullet": True,
            "spacing": True
        },  # Content with bullet point and spacing
        {
            "text": "",
            "bullet": False,
            "spacing": True
        },  # Content with bullet point and no spacing
        {
            "text": f"{trend_observations} is observed in the VRMS trend rate chart over the period",
            "bullet": True,
            "spacing": True
        }  # Content with bullet point and no spacing
    ]

    second_row = table.rows[1]  # Accessing the second row
    cell = second_row.cells[0]  # Accessing the first cell in the second row

    # Add the content to the cell as paragraphs
    for para_content in content:
        p = cell.add_paragraph()
        if isinstance(para_content, dict) and para_content.get("bold"):
            run = p.add_run(para_content["text"])
            run.bold = True  # Bold the text
            run.font.name = 'Calibri'  # Set font style to Calibri (Body)
            run.font.size = Pt(12)  # Set font size
        else:
            run = p.add_run(para_content["text"])
            run.font.name = 'Calibri'  # Set font style to Calibri (Body)
            run.font.size = Pt(12)  # Set font size

        # # style = document.styles.get('List Bullet')
        # if para_content.get("bullet"):
        #     p.style = 'List Bullet 2'  # Apply bullet point style

        p_format = p.paragraph_format
        if para_content.get("spacing"):
            p_format.space_before = Pt(5)  # 8 pt spacing before the paragraph

        # Add indentation for bullet points
        if para_content.get("bullet"):
            run = p.runs[0]
            run.text = " " * 2 + " " * 1 + run.text  # Adding 6 spaces before and 4 spaces after the bullet

    # Save the updated document
    document.save(doc_path)


def Z_Axis(input_df, axes, equipment_id, remark_text, existing_file_path, config_file_path):
    filtered_data = input_df[input_df['Equipment ID'] == equipment_id][['SL No','Equipment ID','Duration', axes, 'Temp (in Degrees)']]
    slno = str(input_df['SL No'])
    Equipment_ID = equipment_id
    Duration = str(filtered_data['Duration'].item())
    Axis_name = str(filtered_data[axes].item())
    Temp = str(filtered_data['Temp (in Degrees)'].item())

    # Add Heading
    doc_path = existing_file_path

    # Load the existing document
    document = Document(doc_path)

    # Get the 'data_Input_folder' from the config
    directory_path = config['data_Input_folder']

    # Create an output folder if it doesn't exist
    output_folder = config['Chart_path']
    # Call the preprocess_data function with the directory path
    dataframes = preprocess_data(directory_path, Equipment_ID)

    dataframes = dataframes[2]

    # Iterate through the dataframes and apply selected trending methods
    plotting(dataframes, output_folder, config)

    Trendline, Moving_Avg = moving_avg_trendline_dataframe(dataframes, config)

    TL_MA_df = pd.DataFrame({'Trend_Line': Trendline, 'Moving_Average': Moving_Avg})

    #Check the health of the equipment(Data:Moving average and TrendLine)
    overall_result = overall_condition(TL_MA_df)

    # Save the updated document
    document.save(doc_path)

    #Velocity RMS Data Summary Table

    # Add a subheading
    sub_heading_text = f"Velocity RMS Data Summary - {Axis_name}:"
    sub_heading_color = (0, 0, 0)  # RGB color values for black
    sub_heading_font = "Calibri (Body)"  # Font name
    font_size = 14
    add_heading_to_document(document, sub_heading_text, 2, sub_heading_color,
                            sub_heading_font, font_size)

    table = document.add_table(rows=2, cols=1, style='Table Grid')


    # Set the first row height to be three times larger than the second row
    for index, row in enumerate(table.rows):
        if index == 0:
            row.height = Cm(10)  # Set the first row height (in inches)
        else:
            row.height = Cm(4)  # Set the second row height (1/3 of the first row)



    # Access the cell at position [0, 0] (first row, first column)
    cell = table.cell(0, 0)

    # Add an image to the cell
    image_path = config['Chart_path']+'\\'+Equipment_ID+'_'+'X_Axis'+'.jpeg' # Replace with the path to your image file
    cell_paragraph = cell.paragraphs[0]
    run = cell_paragraph.add_run()
    run.add_picture(image_path, width=Cm(17.75))  # Adjust width as needed

    moving_average_observations = analyze_condition(overall_result)
    trend_observations = analyze_trend(Trendline)

    content = [
        {"text": "Observations:", "bold": True},  # Bold "Observations"
        {
            "text": moving_average_observations,
            "bullet": True,
            "spacing": True
        },  # Content with bullet point and spacing
        {
            "text": "",
            "bullet": False,
            "spacing": True
        },  # Content with bullet point and no spacing
        {
            "text": f"{trend_observations} is observed in the VRMS trend rate chart over the period",
            "bullet": True,
            "spacing": True
        }  # Content with bullet point and no spacing
    ]

    second_row = table.rows[1]  # Accessing the second row
    cell = second_row.cells[0]  # Accessing the first cell in the second row

    # Add the content to the cell as paragraphs
    for para_content in content:
        p = cell.add_paragraph()
        if isinstance(para_content, dict) and para_content.get("bold"):
            run = p.add_run(para_content["text"])
            run.bold = True  # Bold the text
            run.font.name = 'Calibri'  # Set font style to Calibri (Body)
            run.font.size = Pt(12)  # Set font size
        else:
            run = p.add_run(para_content["text"])
            run.font.name = 'Calibri'  # Set font style to Calibri (Body)
            run.font.size = Pt(12)  # Set font size

        # # style = document.styles.get('List Bullet')
        # if para_content.get("bullet"):
        #     p.style = 'List Bullet 2'  # Apply bullet point style  # need to check , but yet to start at 4/12/23 --> ps

        p_format = p.paragraph_format
        if para_content.get("spacing"):
            p_format.space_before = Pt(5)  # 8 pt spacing before the paragraph

        # Add indentation for bullet points
        if para_content.get("bullet"):
            run = p.runs[0]
            run.text = " " * 2 + " " * 1 + run.text  # Adding 6 spaces before and 4 spaces after the bullet

    #remark wala..
    document.add_paragraph()
    table = document.add_table(rows=1, cols=2, style='Table Grid')
    table.allow_autofit = False
    table.autofit = False
    # table.allow_autofit = True
    # Set individual column widths
    max_width_first_column = Cm(1.91)
    max_width_second_column = Cm(16.46)

    table.columns[0].width = max_width_first_column
    table.columns[1].width = max_width_second_column

    text = "Remark"
    cell = table.cell(0, 0)
    cell.text = text

    cell_1 = table.cell(0, 1)
    cell_1.text = remark_text(overall_result)

    # Calculate the height of the cell based on content
    font_size = 12  # Adjust the font size as needed
    content_height = len(text.split("\n")) * font_size * 1.2  # Adjust the factor as needed

    # Set font style and size for both columns
    for col in [cell, cell_1]:
        run = col.paragraphs[0].runs[0]
        run.font.name = 'Calibri'
        run.font.size = Pt(font_size)

    # Make the text in the first column bold
    run_first_column = cell.paragraphs[0].runs[0]
    run_first_column.bold = True

    cell.height = Cm(content_height)

    # # Save the updated document
    # document.save(doc_path)

        # Save the updated document
    document.save(doc_path)


Axis = ['X-Axis', 'Y-Axis', 'Z-Axis']


# pm modify
for row_num, row in equipment_ids_df.iterrows():
    equipment_id = row[equipment_id_column_name]
    print(equipment_id)

   # prashath's code merge here for plot and observation

        #call x , y, z axis func here # ps

    for axes in Axis:
        if axes == "X-Axis":
            X_Axis(input_df, axes, equipment_id, existing_file_path, config_file_path)
        elif axes == "Y-Axis":
            Y_Axis(input_df, axes, equipment_id, existing_file_path, config_file_path)
        elif axes == "Z-Axis":
            Z_Axis(input_df, axes, equipment_id, remark_text, existing_file_path, config_file_path)
    else:
        print("No data found for Equipment ID:", equipment_id)

###task 18 (adding page number to footer )
###code from supriya


def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)

    # Set font size and style for the page number
    for run in [page_run, page_num_run, of_run, num_pages_run]:
        run.font.size = Pt(12)
        run.font.name = "Lato Heavy"

# Load your existing Word document
doc = Document(existing_file_path)

# Iterate through each section of the document
for section in doc.sections:
    # Check if the section has a footer
    if section.footer is not None:
        # Check if the footer has paragraphs
        if section.footer.paragraphs:
            # Check if the page number is already present in the footer
            if "Page " not in section.footer.paragraphs[0].text:
                # If not present, add the page number to the footer
                add_page_number(section.footer.paragraphs[0])
        else:
            # If the footer does not have paragraphs, create one and add the page number
            paragraph = section.footer.add_paragraph()
            add_page_number(paragraph)
    else:
        # If the section does not have a footer, create one and add the page number
        footer = section.footer
        paragraph = footer.paragraphs[0].add_run()
        add_page_number(paragraph)

# Save the modified document
doc.save(existing_file_path)
# #pip install docx2pdf
from docx2pdf import convert

# Specify the input and output file paths
input_docx_path = "output_report_final.docx"
output_pdf_path = "output_pm.pdf"
convert("output_report_final.docx")



#from supriyatask11
#content code

def extract_page_numbers(pdf_path, search_terms):
    doc = fitz.open(pdf_path)

    page_numbers = {}

    for term in search_terms:
        term = term.lower()
        for page_number in range(doc.page_count):
            page = doc[page_number]
            text = page.get_text("text")
            if term in text.lower():
                page_numbers[term] = page_number + 1
                break

    doc.close()

    return page_numbers


def save_results_to_document(results, output_path):
    doc = Document()

    for term, page_number in results.items():
        if term.isdigit():
            parts = term.split('.')
            if len(parts) > 1:
                indentation = (len(parts) - 1) * '\t'
                content = f"{indentation}{term} {page_number}"
            else:
                content = f"{term} = {page_number}"
        else:
            content = f"{term} = {page_number}"
        doc.add_paragraph(content)

    doc.save(output_path)

'''
def add_colon_to_words(doc_path, terms):
    doc = Document(doc_path)

    for term in terms:
        for paragraph in doc.paragraphs:
            if term.lower() in paragraph.text.lower():
                new_text = paragraph.text.replace(term, f"{term}:")
                paragraph.clear()
                paragraph.add_run(new_text)

    doc.save(doc_path)

'''
def extract_page_numbers_from_content_doc(doc):
    page_numbers = {}

    for paragraph in doc.paragraphs:
        match = re.match(r'(.+): = (\d+)$', paragraph.text.strip())
        if match:
            content = match.group(1).lower()
            page_number = int(match.group(2))
            page_numbers[content] = page_number

    return page_numbers


def add_page_numbers_to_loaded_doc(loaded_doc, page_numbers):
    for paragraph in loaded_doc.paragraphs:
        for content, page_number in page_numbers.items():
            if content.lower() in paragraph.text.lower():
                # Replace the existing page number with the updated page number
                new_text = re.sub(rf'{re.escape(content)} = (\d+)', f'{content} = {page_number}', paragraph.text)
                paragraph.clear()
                paragraph.add_run(new_text)
                break  # Stop searching for this content item once found


def add_table_of_contents(doc, equipment_ids_df, paragraph, page_numbers):
    # Add a line break after the "CONTENTS" heading
    paragraph.add_run('\n')

    # Add the main sections
    sections = [
        ("1. Summary", "Summary"),
        ("2. Equipment Overview Status", "Equipment Overview Status"),
        ("3. Equipment Details", "Equipment Details"),
    ]

    # Set the default font for the entire document
    default_font = doc.styles['Normal'].font
    default_font.name = 'Calibri'
    default_font.size = Pt(18)
    default_font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

    # Initialize the section number
    section_number = 1

    # Iterate through main sections
    for section, section_text in sections:
        # Calculate the number of dots needed to align the page numbers to the right
        dots_needed = 78 - len(section_text) - len(str(page_numbers.get(section_text.lower(), '')))

        # Add the main section text with numbering
        section_with_numbering = f"{section_number}. {section_text}"
        run_section = paragraph.add_run(section_with_numbering)
        run_section.font.size = Pt(18)
        run_section.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

        # Add the dots for alignment
        dots = '.' * dots_needed
        paragraph.add_run(f"{dots}{page_numbers.get(section_text.lower(), '')}")

        # Remove bold formatting
        run_section.bold = False

        # Add a line break after each main section
        paragraph.add_run('\n')

        if section_text.lower() == "equipment details":
            # Add subsections dynamically with generated page numbers only for "Equipment Details"
            for index, row in equipment_ids_df.iterrows():
                target_column_name = 'Equipment ID'
                target_column_lower = target_column_name.lower()

                # Check if the target column name exists in the DataFrame (case insensitive)
                equipment_id_column_name = next(
                    (col_name for col_name in equipment_ids_df.columns if col_name.lower() == target_column_lower),
                    None
                )

                if equipment_id_column_name:
                    equipment_id = row.get(equipment_id_column_name, "")

                    # Trim equipment name if it's too long
                    equipment_name = equipment_id[:100]  # Assuming a maximum length of 100 characters

                    # Construct the subsection without dots
                    subsection = f"   {section_number}.{index + 1}. {equipment_name}"
                    run_subsection = paragraph.add_run(subsection)
                    run_subsection.font.size = Pt(18)
                    run_subsection.font.color.rgb = RGBColor(0, 0, 0)  # Set font color to black

                    # Calculate the number of dots needed to align the page numbers to the right
                    dots_needed_subsection = 60 - len(subsection) - len(str(page_numbers.get(equipment_name.lower(), '')))

                    # Add the dots for alignment
                    dots_subsection = '.' * dots_needed_subsection
                    paragraph.add_run(f"{dots_subsection}{page_numbers.get(equipment_name.lower(), '')}")

                    # Remove bold formatting
                    run_subsection.bold = False

                    # Add a line break after each subsection
                    paragraph.add_run('\n')

        # Increment the section number
        section_number += 1


# Example usage
excel_file = "inputs_for_report.xlsx"
df = pd.read_excel(excel_file)
pdf_path = "output_report_final.pdf" #"content.pdf"
search_words = ["Summary", "Equipment Overview Status", "Equipment Details"]
dynamic_content_items = list(df["Equipment ID"].astype(str))

# Add colon to search words
search_terms_with_colons = [f"{word}:" for word in search_words]

# Add colon to dynamic content items
dynamic_content_terms_with_colons = [f"{item}:" for item in dynamic_content_items]

# Extract page numbers from the PDF
result = extract_page_numbers(pdf_path, search_terms_with_colons + dynamic_content_terms_with_colons)

print(result)

# Save the results to a single document
output_path = "contentstarts.docx"
save_results_to_document(result, output_path)

# Add colon to search words and dynamic content items
# add_colon_to_words(output_path, search_terms_with_colons + dynamic_content_terms_with_colons)

# Load the existing Word document
loaded_doc = Document(existing_file_path)#"content.docx")

# Assume that the "CONTENTS" heading is present in the document
# If not, you may need to modify this part based on your document structure78

for i, paragraph in enumerate(loaded_doc.paragraphs):
    if "CONTENTS" in paragraph.text:
        excel_file_path = 'inputs_for_report.xlsx'  # Replace with your actual Excel file path
        equipment_ids_df = pd.read_excel(excel_file_path)

        # Load the document with extracted page numbers
        content_doc = Document("contentstarts.docx")

        # Extract page numbers from the content document
        page_numbers = extract_page_numbers_from_content_doc(content_doc)

        # Add table of contents with page numbers
        add_table_of_contents(loaded_doc, equipment_ids_df, paragraph, page_numbers)
        break  # Stop searching once found

# Save the modified document
loaded_doc.save("updated_contents_with_page_numbers_final_report.docx")
print("hooo")






#pm add content page format (...)
def format_contents_page(doc):
    # Find the "CONTENTS" heading
    for i, paragraph in enumerate(doc.paragraphs):
        if "CONTENTS" in paragraph.text:
            contents_heading_index = i
            break
    else:
        # Handle case where "CONTENTS" heading is not found
        print("CONTENTS heading not found.")
        return

    # Iterate through paragraphs starting from the "CONTENTS" heading
    for paragraph in doc.paragraphs[contents_heading_index + 1:]:
        # Break the loop if an empty paragraph is encountered
        if not paragraph.text.strip():
            break

        # Split the paragraph into section title and page number
        parts = paragraph.text.split('......')
        if len(parts) == 2:
            section_title, page_number = parts
        else:
            # Handle cases where the format is different
            print(f"Could not parse line: {paragraph.text}")
            continue

        # Calculate the indentation for right-aligning the page number
        indentation = 80 - len(section_title) - len(page_number)

        # Format the line with left-aligned section title and right-aligned page number
        formatted_line = f"{section_title.strip()}: {'.' * indentation}{page_number.strip()}"

        # Update the paragraph text with the formatted line
        paragraph.text = formatted_line
# Example usage
existing_file_path_with_content = "updated_contents_with_page_numbers_final_report.docx"
loaded_doc = Document(existing_file_path_with_content)

# Format the "CONTENTS" page
format_contents_page(loaded_doc)

# Save the modified document
loaded_doc.save("trends_report_final.docx")

print("omg")

convert("trends_report_final.docx")

print("jjjj")