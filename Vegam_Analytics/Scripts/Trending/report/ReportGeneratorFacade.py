# ReportGeneratorFacade.py
from datetime import date, timedelta, datetime
import pytz
from datetime import datetime

# helps to communicate with other  4 classes and get report
# run the file, get output --> this is goal
import errno
from docx.shared import Pt
from docx import Document
from Scripts.Trending.report.ExtractMetaData import ExtractMetaData
from Scripts.Trending.report.ExtractSensorData import ExtractSensorData
from Scripts.Trending.report.AverageCalculator import AverageCalculator
# from Scripts.Trending.report.ReportGenerator import ReportGenerator
from Scripts.Trending.report.DailyReportGenerator import DailyReportGenerator
import json
from datetime import datetime
import time
import os
import logging
from logging.handlers import RotatingFileHandler
import pytz
import psutil
from docx2pdf import convert
import pandas as pd
import sys



log_folder_path = 'C:\\Vegam\\Trends_alert\\logs'
# log_folder_path = 'logs'
if not os.path.exists(log_folder_path):
    try:
        os.makedirs(log_folder_path)
        logging.info(f'Created "{log_folder_path}" folder.')
    except Exception as e:
        logging.error(f'Error creating "{log_folder_path}" folder: {str(e)}')

logger = logging.getLogger()
log_format = '%(asctime)s - [%(filename)s::%(lineno)d] - %(levelname)12s - %(threadName)22s  - %(funcName)s -' \
             ' %(message)s'
logging.basicConfig(
    format=log_format,
    datefmt='%Y-%m-%d %H:%M:%S',
    level=logging.DEBUG
)

file_handler = RotatingFileHandler(
    'C:\\Vegam\\Trends_alert\\logs\\daily_report.log', maxBytes=200 * 1024 * 1024, backupCount=10, mode='a'
)
# file_handler = RotatingFileHandler(
#     'logs/daily_report.log', maxBytes=200 * 1024 * 1024, backupCount=10, mode='a'
# )
file_handler.setFormatter(logging.Formatter(log_format))
logger.addHandler(file_handler)



class ReportGeneratorFacade:
    def __init__(self, file_path, area,starting_date):
        try:
            self.filepath = file_path  # config_filename
            with open(self.filepath, 'r') as json_file:
                self.Trends_Config = json.load(json_file)
            self.cache_time = 1  # time in seconds
            self.report_start_hour = self.Trends_Config["trends"]["report_start_hour"]
            self.report_start_minute = self.Trends_Config["trends"]["report_start_minute"]
            self.report_start_second = self.Trends_Config["trends"]["report_start_second"]
            self.report_start_microsecond = self.Trends_Config["trends"]["report_start_microsecond"]

            self.report_end_hour = self.Trends_Config["trends"]["report_end_hour"]
            self.report_end_minute = self.Trends_Config["trends"]["report_end_minute"]
            self.report_end_second = self.Trends_Config["trends"]["report_end_second"]
            self.report_end_microsecond = self.Trends_Config["trends"]["report_end_microsecond"]
            self.report_days_offset = self.Trends_Config["trends"]["report_days_offset"]

            # self.exclude_mac_ids = self.Trends_Config["trends"]["exclude_mac_ids"]

            # self.area_check
            # self.warning_value_for_report_generator = self.Trends_Config["trend_generate"]["Warning"]
            # self.critical_value_for_report_generator = self.Trends_Config["trend_generate"]["Critical"]

            self.machine_imange_unavailable_pic_path = self.Trends_Config["trends"]["machine_imange_unavailable_pic_path"]
            #self.area_check = self.Trends_Config["trends"]["area_check"]
            self.area_check = area
            self.exclude_mac_ids = self.Trends_Config["trends"]["exclude_mac_ids"].get(self.area_check, [])

            self.machine_pic_path = self.Trends_Config["trends"]["machine_image_path"]
            self.moving_avg_plot_path = self.Trends_Config["trends"]["Moving_Average_plots_Path"]
            self.metadata_extractor = ExtractMetaData(self.Trends_Config["trends"]["vmaint_api"], self.cache_time,
                                                      self.area_check)
            self.sensor_data_extract = ExtractSensorData(self.Trends_Config["trends"]["vegam_view_prod_api"])
            self.outputs_folder_path = os.path.join(self.Trends_Config["trends"]["output_report_path"], starting_date)

            # Ensure the folder for the starting_date exists; create it if not
            os.makedirs(self.outputs_folder_path, exist_ok=True)

            # Construct the folder path for the current area
            self.area_folder_path = os.path.join(self.outputs_folder_path, self.area_check)
            #logger.info(self.area_folder_path,'1111111111111111111111111111111111111111')
            # Ensure the area-specific folder exists; create it if not
            os.makedirs(self.area_folder_path, exist_ok=True)
            # Construct the paths for report and PDF based on the area-specific folder
            new_doc_file_name = f'{self.area_check}.docx'
            self.existing_file_path= os.path.join(self.area_folder_path, new_doc_file_name)########report_existing_path
            #print(self.existing_file_path,'222222222222222222222222222')

            new_pdf_file_name = f'{self.area_check}.pdf'
            self.pdf_path = os.path.join(self.area_folder_path, new_pdf_file_name)##########report pdf path
            #print(self.pdf_path,'333333333333333333')

            self.list_of_tags_macids = []
            self.warning_value_for_report_generator = self.Trends_Config["trend_generate"][self.area_check]["Warning"]
            self.critical_value_for_report_generator = self.Trends_Config["trend_generate"][self.area_check]["Critical"]

            # self.existing_file_path = self.Trends_Config["trends"]["report_existing_path"]
            # self.pdf_path = self.Trends_Config["trends"]["pdf_path"]
            # self.report_existing_path = self.Trends_Config["trends"]["report_existing_path"]
            # new_doc_file_name = f'{self.area_check}.docx'
            # print(new_doc_file_name)
            # self.report_existing_path = self.report_existing_path +"/"+new_doc_file_name
            #
            # print(self.report_existing_path,'222222222222222222222')
            # new_pdf_file_name = f'{self.area_check}.pdf'
            # print(new_pdf_file_name)
            #
            # self.report_pdf_path = self.report_existing_path +"/"+new_pdf_file_name
            # print(self.report_pdf_path,'3333333333333333333')
            #


            #self.report_pdf_path = f"C:\\Vegam\\Trends_alert\\test_output\\report_{self.area_check}_{self.starting_time}.pdf"

            self.original_file_path = self.Trends_Config["trends"]["report_input_template"]
            self.vegam_logo_path = self.Trends_Config["trends"]["report_input_vegam_logo"]
            # self.outputs_folder_path = self.Trends_Config["trends"]["output_report_path"]
            #self.st = self.Trends_Config["trends"]["start_time_report"]
            #self.et = self.Trends_Config["trends"]["end_time_report"]
            #self.trends_json_path = self.Trends_Config["alert_payload"]["trend_path"]
            self.alert_json_path = self.Trends_Config["alert_payload"]["alert_path"]
            self.trends_date = self.Trends_Config["alert_payload"]["toDate"]
            self.list_of_equips = []
            self.filter_sensor_list_with_tag = []
            self.list_of_equips_with_mac_id = []
            self.count = 1
            self.percentage_threshold = self.Trends_Config["trends"]["percentage_threshold"]

            self.report_generate_obj = None
            self.meta_data = self.get_metadata_json_from_vamint_api_for_report()
            if self.meta_data:
                # print("entered meta_data")
                self.customer_name = self.meta_data["SiteName"]
                self.get_euipments_details_with_mac_id_for_summary_status()  # mac id count we can check [ not much needed this func]
                if self.list_of_equips_with_mac_id:
                    # print("entered list_of_equips_with_mac_id")
                    self.get_euipments_details_for_summary()
                    if self.list_of_equips:
                        # print("entered list_of_equips")
                        self.report_generate_obj = DailyReportGenerator(self.existing_file_path,self.pdf_path, self.original_file_path,
                                                                   self.vegam_logo_path, self.customer_name,
                                                                   self.list_of_equips,self.list_of_equips_with_mac_id ,self.outputs_folder_path,
                                                                   self.area_check,self.critical_value_for_report_generator,self.warning_value_for_report_generator,self.report_start_hour,self.report_start_minute,self.report_start_second,self.report_start_microsecond
                                                                   ,self.report_end_hour,self.report_end_minute,self.report_end_second,self.report_end_microsecond,self.report_days_offset,self.percentage_threshold )
            self.doc = None  # Document(self.existing_file_path)

            self.obs_remarks_status_dict = {}
            #self.input_dict = {}
            self.axis_values = {}

        except Exception as w:
            logger.error(f"An error occurred during initialization in ReportGeneratorFacade: {str(w)}")

    def get_metadata_from_vamint_api_for_report_list_of_sensor_id_with_tag(
            self):
        try:
            # PURPOSE : getting all mac ids[ every area]
            list_of_tags_with_macids = self.metadata_extractor.get_mac_ids_tags_for_report_with_temp()
            self.list_of_tags_macids = list_of_tags_with_macids
            # return list_of_tags_with_macids

        except Exception as e:
            logger.error(f"Error in get_metadata_from_vamint_api_for_report_list_of_sensor_id_with_tag: {e}")

    def get_metadata_json_from_vamint_api_for_report(self):

        # logger.info(
        #     f"start of meta data getting method in thread..----   {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        try:
            # PURPOSE : get meta data json..
            meatadata_json = self.metadata_extractor.get_metadata_from_api()
            return meatadata_json

        except Exception as e:
            logger.error(f"Error in get_metadata_json_from_vamint_api_for_report: {e}")


    def make_report(self):
        try:
            #input_dict_list = self.get_input_dict()
            #print(input_dict_list,'input_dict_list from get_input_dict method ')
            input_dict_list = self.get_save_vrms_data_to_get_report()
            logger.info(f"{input_dict_list} input_dict")
            if input_dict_list:
                # input_dict_list = [{'MAC_ID': 'C258249CFE98', 'Velocity RMS - Radial Horizontal': 1.045885202390778,
                #   'Velocity RMS - Radial Vertical': 1.9041070622570058, 'Velocity RMS - Axial': 2.376408072229036,
                #   'Equipment ID': '1200MW_CT_U1_Fan1'},
                #  {'MAC_ID': 'D7E0DB929BF9', 'Velocity RMS - Radial Horizontal': 2.65018826385067,
                #   'Velocity RMS - Radial Vertical': 6.215167415258651, 'Velocity RMS - Axial': 4.597339692331943,
                #   'Equipment ID': '1200MW_CT_U4_Fan9'}]

                logger.info('hi,i got final inputdict')
                # self.doc.save(self.existing_file_path)
                # convert(self.existing_file_path)

                get_equipid_and_macid = self.filter_mac_ids_with_tag()
                logger.info(f"{get_equipid_and_macid} get_equipid_and_macid")

                # self.report_generate_obj.create_report_common(input_dict_list)
                # print('hi,table created')

                result_list = []
                # self.get_save_vrms_data_to_get_report()
                axis_info = self.metadata_extractor.get_axis_information()
                logger.info(f"{axis_info} - axisssssssssssssss")##{'D248811B8F5C': [{'Topic': '3426', 'TopicRepresents': 'V_RMSX', 'TopicType': 'S'}, {'Topic': '3431', 'TopicRepresents': 'V_RMSY', 'TopicType': 'S'}, {'Topic': '3436', 'TopicRepresents': 'V_RMSZ', 'TopicType': 'S'}, {'AxisX': 'VERTICAL_R', 'AxisY': 'AXIAL', 'AxisZ': 'HORIZONTAL_R'}, {'macid': 'D248811B8F5C'}]}, {'ECCCA56F4C24': [{'Topic': '3454', 'TopicRepresents': 'V_RMSX', 'TopicType': 'S'}, {'Topic': '3459', 'TopicRepresents': 'V_RMSY', 'TopicType': 'S'}, {'Topic': '3464', 'TopicRepresents': 'V_RMSZ', 'TopicType': 'S'}, {'AxisX': 'VERTICAL_R', 'AxisY': 'AXIAL', 'AxisZ': 'HORIZONTAL_R'}, {'macid': 'ECCCA56F4C24'}]}, {'C852DF362F8A': [{'Topic': '3485', 'TopicRepresents': 'V_RMSX', 'TopicType': 'S'}, {'Topic': '3490', 'TopicRepresents': 'V_RMSY', 'TopicType': 'S'}, {'Topic': '3495', 'TopicRepresents': 'V_RMSZ', 'TopicType': 'S'}, {'AxisX': 'VERTICAL_R', 'AxisY': 'AXIAL', 'AxisZ': 'HORIZONTAL_R'}, {'macid': 'C852DF362F8A'}]}, {'C3F327CE864E': [{'Topic': '3513', 'TopicRepresents': 'V_RMSX', 'TopicType': 'S'}, {'Topic': '3518', 'TopicRepresents': 'V_RMSY', 'TopicType': 'S'}, {'Topic': '3523', 'TopicRepresents': 'V_RMSZ', 'TopicType': 'S'}, {'AxisX': 'VERTICAL_R', 'AxisY': 'AXIAL', 'AxisZ': 'HORIZONTAL_R'}, {'macid': 'C3F327CE864E'}]}, {'C6BCBB209A02': [{'Topic': '3544', 'TopicRepresents': 'V_RMSX', 'TopicType': 'S'}, {'Topic': '3549', 'TopicRepresents': 'V_RMSY', 'TopicType': 'S'}, {'Topic': '3554', 'TopicRepresents': 'V_RMSZ', 'TopicType': 'S'}, {'AxisX': 'VERTICAL_R', 'AxisY': 'AXIAL', 'AxisZ': 'HORIZONTAL_R'}, {'macid': 'C6BCBB209A02'}]}, {'EB9477C86B7B': [{'Topic': '3572', 'TopicRepresents': 'V_RMSX', 'TopicType': 'S'}, {'Topic': '3577', 'TopicRepresents': 'V_RMSY', 'TopicType': 'S'}, {'Topic': '3582', 'TopicRepresents': 'V_RMSZ', 'TopicType': 'S'}, {'AxisX': 'VERTICAL_R', 'AxisY': 'AXIAL', 'AxisZ': 'HORIZONTAL_R'}, {'macid': 'EB9477C86B7B'}]}, {'C86D499BDE9A': [{'Topic': '3603', 'TopicRepresents': 'V_RMSX', 'TopicType': 'S'}, {'Topic': '3608', 'TopicRepresents': 'V_RMSY', 'TopicType': 'S'}, {'Topic': '3613', 'TopicRepresents': 'V_RMSZ', 'TopicType': 'S'}, {'AxisX': 'VERTICAL_R', 'AxisY': 'AXIAL', 'AxisZ': 'HORIZONTAL_R'}, {'macid': 'C86D499BDE9A'}]}, {'D02E602AE46B': [{'Topic': '3631', 'TopicRepresents': 'V_RMSX', 'TopicType': 'S'}, {'Topic': '3636', 'TopicRepresents': 'V_RMSY', 'TopicType': 'S'}, {'Topic': '3641', 'TopicRepresents': 'V_RMSZ', 'TopicType': 'S'}, {'AxisX': 'VERTICAL_R', 'AxisY': 'AXIAL', 'AxisZ': 'HORIZONTAL_R'}, {'macid': 'D02E602AE46B'}]}, {'D373CFE3A391': [{'Topic': '3662', 'TopicRepresents': 'V_RMSX', 'TopicType': 'S'}, {'Topic': '3667', 'TopicRepresents': 'V_RMSY', 'TopicType': 'S'}, {'Topic': '3672', 'TopicRepresents': 'V_RMSZ', 'TopicType': 'S'}, {'AxisX': 'VERTICAL_R', 'AxisY': 'AXIAL', 'AxisZ': 'HORIZONTAL_R'}, {'macid': 'D373CFE3A391'}]}, {'F9DA7C304C97': [{'Topic': '3690', 'TopicRepresents': 'V_RMSX', 'TopicType': 'S'}, {'Topic': '3695', 'TopicRepresents': 'V_RMSY', 'TopicType': 'S'}, {'Topic': '3700', 'TopicRepresents': 'V_RMSZ', 'TopicType': 'S'}, {'AxisX': 'VERTICAL_R', 'AxisY': 'AXIAL', 'AxisZ': 'HORIZONTAL_R'}, {'macid': 'F9DA7C304C97'}]}, {'DB316051FCFE': [{'Topic': '3864', 'TopicRepresents': 'V_RMSX', 'TopicType': 'S'}, {'Topic': '3869', 'TopicRepresents': 'V_RMSY', 'TopicType': 'S'}, {'Topic': '3874', 'TopicRepresents': 'V_RMSZ', 'TopicType': 'S'}, {'AxisX': 'VERTICAL_R', 'AxisY': 'AXIAL', 'AxisZ': 'HORIZONTAL_R'}, {'macid': 'DB316051FCFE'}]}, {'D2930832E760': [{'Topic': '3892', 'TopicRepresents': 'V_RMSX', 'TopicType': 'S'}, {'Topic': '3897', 'TopicRepresents': 'V_RMSY', 'TopicType': 'S'}, {'Topic': '3902', 'TopicRepresents': 'V_RMSZ', 'TopicType': 'S'}, {'AxisX': 'VERTICAL_R', 'AxisY': 'AXIAL', 'AxisZ': 'HORIZONTAL_R'}, {'macid': 'D2930832E760'}]}]


                for equipid_dict in get_equipid_and_macid:
                    equipid_mac_id = equipid_dict.get('MAC_ID', None)
                    for input_dict in input_dict_list:
                        input_mac_id = input_dict.get('MAC_ID', None)

                        # Find the corresponding axis_info for the current input_mac_id
                        # matching_axis_info = next(
                        #     (axis_info[input_mac_id] for axis_info_key in axis_info if axis_info_key == input_mac_id),
                        #     None)
                        matching_axis_info = next(
                            (axis_info_key[input_mac_id] for axis_info_key in axis_info if list(axis_info_key.keys())[0] == input_mac_id),
                            None)

                        if matching_axis_info:
                            # Extract values for 'AxisX', 'AxisY', and 'AxisZ' if available
                            axis_x_value = next((item['AxisX'] for item in matching_axis_info if 'AxisX' in item), None)
                            axis_y_value = next((item['AxisY'] for item in matching_axis_info if 'AxisY' in item), None)
                            axis_z_value = next((item['AxisZ'] for item in matching_axis_info if 'AxisZ' in item), None)

                            # Iterate over input_dict_list
                            # for input_dict in input_dict_list:
                            #     input_mac_id = input_dict.get('MAC_ID', None)


                            if input_mac_id == equipid_mac_id:
                                # Use these values as needed in your logic
                                logger.info(
                                    f"For MAC_ID {input_mac_id}: AxisX = {axis_x_value}, AxisY = {axis_y_value}, AxisZ = {axis_z_value}")

                                # Example mapping logic:
                                mapping_result_dict = {
                                    'EquipmentName_MacID': equipid_dict.get('EquipmentName_MacID', None),
                                    f'Velocity RMS - {axis_x_value}': input_dict.get('X', None),
                                    f'Velocity RMS - {axis_y_value}': input_dict.get('Y', None),
                                    f'Velocity RMS - {axis_z_value}': input_dict.get('Z', None),
                                    f'Velocity RMS - {axis_x_value}_% Data Beyond Critical': input_dict.get(
                                        'X_Percentage_Beyond_Critical', None),
                                    f'Velocity RMS - {axis_y_value}_% Data Beyond Critical': input_dict.get(
                                        'Y_Percentage_Beyond_Critical', None),
                                    f'Velocity RMS - {axis_z_value}_% Data Beyond Critical': input_dict.get(
                                        'Z_Percentage_Beyond_Critical', None),
                                }

                                replacement_mapping = {
                                    'Velocity RMS - HORIZONTAL_R': 'Velocity RMS - Radial Horizontal',
                                    'Velocity RMS - VERTICAL_R': 'Velocity RMS - Radial Vertical',
                                    'Velocity RMS - AXIAL': 'Velocity RMS - Axial',
                                    'Velocity RMS - HORIZONTAL_R_% Data Beyond Critical': 'Velocity RMS - Radial Horizontal_% Data Beyond Critical',
                                    'Velocity RMS - VERTICAL_R_% Data Beyond Critical': 'Velocity RMS - Radial Vertical_% Data Beyond Critical',
                                    'Velocity RMS - AXIAL_% Data Beyond Critical': 'Velocity RMS - Axial_% Data Beyond Critical'

                                }
                                mapping_result_dict = {replacement_mapping.get(key, key): value for key, value in
                                                       mapping_result_dict.items()}

                                result_list.append(mapping_result_dict)
                                logger.info(f"{result_list} - final result list")
                            else:
                                logger.info("Equipment MacID not matched")
                        else:
                            logger.info(f"No matching axis_info found for MAC_ID {equipid_mac_id}")
                    # if input_mac_id == equipid_mac_id:
                    #     combined_dict = {
                    #         'EquipmentName_MacID': equipid_dict.get('EquipmentName_MacID', None),
                    #         'Velocity RMS - Radial Horizontal': input_dict.get('X', None),
                    #         'Velocity RMS - Radial Vertical': input_dict.get('Y', None),
                    #         'Velocity RMS - Axial': input_dict.get('Z', None),
                    #         'Velocity RMS - Radial Horizontal_% Data Beyond Critical': input_dict.get(
                    #             'X_Percentage_Beyond_Critical', None),
                    #         'Velocity RMS - Radial Vertical_% Data Beyond Critical': input_dict.get(
                    #             'Y_Percentage_Beyond_Critical', None),
                    #         'Velocity RMS - Axial_% Data Beyond Critical': input_dict.get(
                    #             'Z_Percentage_Beyond_Critical', None)
                    #
                    #     }
                    #     result_list.append(combined_dict)
                logging.info(f"{result_list} result_list")


                self.report_generate_obj.create_report_common(result_list)
                logger.info('hi,table created and values added to cells')

                self.doc = Document(self.existing_file_path)
                self.report_generate_obj.content_page_works(self.doc)

                self.doc.save(self.existing_file_path)
                convert(self.existing_file_path)
                #self.doc.close()
                #self.doc.quit()
            else:
                logger.info(f"input_dict {input_dict_list}")
                logger.info("Empty input_dict_list, no data to store. Doc not generated")


        except Exception as e:
            logger.error(f"Error in make_report: {e}")

    def filter_mac_ids_with_tag(self):
        try:
            # PURPOSE: got every area mac id with tags (vmrs , time, axials) : var : self.list_of_tags_macids
            self.get_metadata_from_vamint_api_for_report_list_of_sensor_id_with_tag()
            # PURPOSE: got mac ids for selected area with equipment details
            filtered_mac_id_list_with_euqip_details = self.metadata_extractor.get_hierarchy_level_with_equipment_details_from_metadata()

            equipment_names = [item['EquipmentName'] for item in filtered_mac_id_list_with_euqip_details]
            # PURPOSE: Sort 'EquipmentName' alphabetically to show in report prettily
            sorted_equipment_names = sorted(equipment_names)

            # PURPOSE: Rearrange the list of mac id with equipments based on sorted 'EquipmentName' list
            # purpose -> sorted_data_list is having only "EquipmentName" but sorted
            sorted_data_list = sorted(filtered_mac_id_list_with_euqip_details,
                                      key=lambda x: sorted_equipment_names.index(x['EquipmentName']))

            unfilter_sensor_list_with_tag = self.list_of_tags_macids #purpose: this is every area mac id with tags
            # print(unfilter_sensor_list_with_tag)
            #purpose -> var: filtered_sensor_list_with_equips_sensor_mac_ids will get filtered mac ids only ,
            # - but sorted alphabetically by help of :sorted_data_list
            # sorted data list is like :
            # [{'SensorMacID': 'DB316051FCFE', 'EquipmentID': '8112', 'EquipmentName': 'BC1 TT Conveyor', 'Area_Hierarchy': 'BALCO_Common_CHP'}
            # {'SensorMacID': 'D2930832E760', 'EquipmentID': '8112', 'EquipmentName': 'BC1 TT Conveyor', 'Area_Hierarchy': 'BALCO_Common_CHP'}]
            # filtered_sensor_list_with_equips_sensor_mac_ids is like ['DB316051FCFE', 'D2930832E760']'''
            filtered_sensor_list_with_equips_sensor_mac_ids = []
            for equipment_list in sorted_data_list:  ##########
                # print(equipment_list)
                filtered_sensor_list_with_equips_sensor_mac_ids.append(equipment_list['SensorMacID'])

            # purpose: filter_sensor_list_with_tag will have filtered sorted macids which will go to the for loop to save data and process , so on...
            for a in filtered_sensor_list_with_equips_sensor_mac_ids:
                for item in unfilter_sensor_list_with_tag:
                    # print(item)
                    if a in item:
                        self.filter_sensor_list_with_tag.append(item)  # this is going to for loop ...

            # print(len(self.filter_sensor_list_with_tag),"-----")
            logger.info(f"{len(sorted_data_list)} -----is this two number matching-------")
            transformed_data_list = [{'MAC_ID': item['SensorMacID'], 'EquipmentName': item['EquipmentName'],
                                      'EquipmentName_MacID': f"{item['EquipmentName']}-{item['SensorMacID']}"} for item
                                     in sorted_data_list]

            # print("transformed_data_list - ", transformed_data_list)
            return transformed_data_list

        except Exception as d:
            logger.error(f"Error in get_metadata_json_from_vamint_api_for_report: {d}")
    # def get_equipid_and_macid(self):
    #     data = self.metadata_extractor.get_hierarchy_level_with_equipment_details_from_metadata()
    #     result_list = []
    #
    #     for item in data:
    #         result_dict = {'MAC_ID': item['SensorMacID'], 'Equipment ID': item['EquipmentName']}
    #         result_list.append(result_dict)
    #
    #     print(result_list)
    #     return result_list

    def get_save_vrms_data_to_get_report(self):
        try:
            dict = []
            result_list = []


            end_time_api_time = datetime.now().replace(hour=self.report_end_hour, minute=self.report_end_minute, second=self.report_end_second, microsecond=self.report_end_microsecond)
            start_time_api = datetime.now() - timedelta(days=self.report_days_offset)
            start_time_api_time = start_time_api.replace(hour=self.report_start_hour, minute=self.report_start_minute, second=self.report_start_second, microsecond=self.report_start_microsecond)

            # end_time_api_time_utc1 = end_time_api_time_ist.astimezone(utc1_tz)
            # start_time_api_time_utc1 = start_time_api_time_ist.astimezone(utc1_tz)
            logger.info(f"before timestamp - {start_time_api_time} - {end_time_api_time}")
            start_time = int(start_time_api_time.timestamp() * 1000)
            end_time = int(end_time_api_time.timestamp() * 1000)

            # # global current_time #pm test
            # start_time = int(datetime.strptime(self.Trends_Config["trends"]['start_time_report'],
            #                                    "%d/%m/%y %I:%M %p").timestamp() * 1000)
            # end_time = int(datetime.strptime(self.Trends_Config["trends"]['end_time_report'],
            #                                  "%d/%m/%y %I:%M %p").timestamp() * 1000)
            logger.info(f"after time stamp - {start_time}, {end_time}")
            # start_time =1704043800000# int(datetime.strptime(self.Trends_Config["trends"]['start_time_report'], "%d/%m/%y %I:%M %p").timestamp() * 1000)
            # end_time =1704173400000# int(datetime.strptime(self.Trends_Config["trends"]['end_time_report'], "%d/%m/%y %I:%M %p").timestamp() * 1000)
            # cycle_start_time = int(time.time())
            # Get initial memory usage
            # initial_memory = psutil.virtual_memory().percent
            # initial_memory_usage = psutil.virtual_memory().used
            self.filter_mac_ids_with_tag()  # to get filtered mac id with tags



            if os.path.exists(self.existing_file_path):
                self.doc = Document(self.existing_file_path)
            else:
                # If the file is not found, create a new document
                logger.info(f"File not found at: {self.existing_file_path}")
                self.doc = Document()
                self.doc.save(self.existing_file_path)


            input_data = self.filter_sensor_list_with_tag  # om og
            logger.info(f"{input_data} - mydata hereeeeeeeeeee")
            logger.info(f"this is input data to the for loop.... {len(input_data)}")
            # exit()
            current_time_before_process = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            logger.info(
                f"START of the CYCLE - getting vrms data and processing sensors: current time is : {current_time_before_process}")

            if input_data:
                for sensor_details in input_data:  # what if input data is empty --????? have if else -- pm
                    # print(sensor_details)
                    key = next(iter(sensor_details.keys()))
                    mac_id = key  # "F44F92D7B58C"# key
                    # print(key)

                    # topic_values = [item['Topic'] for item in sensor_details[key]]
                    topic_values = [item['Topic'] for item in sensor_details[key] if 'Topic' in item]
                    self.axis_values = next((item for item in sensor_details[key] if 'AxisX' in item), {})
                    # print(topic_values)
                    # topic_values.append(topic_values.pop(0))
                    # print(topic_values,'topic')
                    # Order the signal IDs as 'vrmsx', 'vrmsy', 'vrmsz'
                    ordered_topic_values = sorted(topic_values,
                                                  key=lambda x: {'vrmsx': 0, 'vrmsy': 1, 'vrmsz': 2}.get(x, 3))

                    logger.info(f"{str(start_time)} - {str(end_time)} ----this is start time , end time in unix milli seconds")
                    payload = {
                        "historyDataParamter": [
                            {"signalIds": ordered_topic_values, "fromTime": start_time, "toTime": end_time}
                        ]
                    }
                    # logger.info("Start getting data for sensor mac id:"+str(mac_id))
                    logger.info(f"{payload}")
                    logger.info(
                        f"start of getting sensor data for  {mac_id} :----   {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                    sensor_data_start_time = int(time.time())
                    vrms_temp_data = self.sensor_data_extract.fetch_vrms_temp_data(payload)
                    logger.info(f"Fetched VRMS Data for Signal IDs {topic_values}:")
                    sensor_data_end_time = int(time.time())
                    fetch_time = sensor_data_end_time - sensor_data_start_time
                    logger.info(f'Data fetching time for {mac_id} is {fetch_time} seconds')

                    # Extract v and t data
                    logger.info(f"len of vrms data from vegam view: {str(len(vrms_temp_data))}")
                    if len(vrms_temp_data) > 0 and mac_id not in self.exclude_mac_ids:

                        extracted_data = self.sensor_data_extract.extract_v_and_t_from_api_data(vrms_temp_data)

                        sheet_name_prefix = f'Sheet'
                        # output_file_prefix = f'C:/Trends_services/outputs/vrms_data'  #saving in output folder
                        output_file_prefix = self.area_folder_path+'/vrms_data'#f'C:/Vegam/Trends_alert/outputs/vrms_data'  # saving in output folder
                        self.make_sure_path_exists(output_file_prefix)
                        self.sensor_data_extract.save_data_to_excel_for_alert(extracted_data, output_file_prefix,
                                                                              sheet_name_prefix, mac_id)

                        # get_equipid_and_macid = self.filter_mac_ids_with_tag()
                        # print(get_equipid_and_macid,'lokijuhygtfresdfgthyjukl;')



                        trends_obj = AverageCalculator(self.list_of_equips,mac_id,self.critical_value_for_report_generator,self.area_folder_path)

                        input_dict  = trends_obj.calculate_average_axis_wise(mac_id)

                        # print('beforeeeeeeeeeeee - ',input_dict)
                        # columns_to_replace = ['X', 'Y', 'Z']

                        # Iterate through the columns and update NaN values
                        for key, value in input_dict.items():
                            if pd.isna(value):
                                input_dict[key] = '-'

                        # Print the updated dictionary
                        # print(input_dict)

                        # print(len(mac_id),'lenmacid')

                        # print('afterrrrrrrrrrrrrrrrrr - ',input_dict)
                        #print(f"Retrieved Equipment ID: {input_dict.get('Equipment ID')}")

                        dict.append(input_dict)
                        logger.info(f"{dict} al dict")


                        # second_list = dict
                        # first_list = get_equipid_and_macid
                        #
                        # print(first_list, '1')
                        # print(second_list, '2')
                        #
                        # for first_dict in first_list:
                        #     for second_dict in second_list:
                        #
                        #         if first_dict['MAC_ID'] == second_dict['MAC_ID']:
                        #             # Replace keys in second_dict
                        #             second_dict_replaced = {
                        #                 'EquipmentName_MacID':first_dict.get('MacID_EquipmentName',None),
                        #                 'Velocity RMS - Radial Horizontal': second_dict.get('X', None),
                        #                 'Velocity RMS - Radial Vertical': second_dict.get('Y', None),
                        #                 'Velocity RMS - Axial': second_dict.get('Z', None),
                        #                 'Velocity RMS - Radial Horizontal_% Data Beyond Critical':second_dict.get('X_Percentage_Beyond_Critical',None),
                        #                 'Velocity RMS - Radial Vertical_% Data Beyond Critical': second_dict.get('Y_Percentage_Beyond_Critical', None),
                        #                 'Velocity RMS - Axial_% Data Beyond Critical': second_dict.get('Z_Percentage_Beyond_Critical', None)
                        #
                        #             }
                        #
                        #             combined_dict = {**first_dict, **second_dict_replaced}
                        #             result_list.append(combined_dict)
                        #
                        # print(result_list,'result_listttttttt')

                        # for first_dict in first_list:
                        #     for second_dict in second_list:
                        #         if first_dict['MAC_ID'] == second_dict['MAC_ID']:
                        #             combined_dict = {**first_dict, **second_dict}
                        #             result_list.append(combined_dict)
                        #
                        # print(result_list,'/////////////////////////////')
                        #



                        self.mac_id_with_equip = next((item for item in self.list_of_equips_with_mac_id if mac_id in item),
                                                      "")

                        if len(self.mac_id_with_equip) > 0:
                            logger.info(f"{self.mac_id_with_equip}")
                        else:
                            logger.info("Mac ID not found in the list.")


                        logger.info(
                            f"End of processing and publishing for   {mac_id}----   {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                        self.previous_end_time = end_time  # save the end_time for the next execution
                        # print(f"All separated values saved successfully.")
                        logger.info("All separated values saved successfully.")
                        self.count = self.count + 1

                    else:
                        logger.info(
                            f"VEGAM VIEW API is returing empty result for sensor: {mac_id}, payload is:{payload} and current time is:  {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
                    # current_time_after_process =datetime.now().strftime('%Y-%m-%d %H:%M:%S')



                # # HACK: deleteing empty pages .. as template which contains only one page is not working ..
                # #  so taking 4 pages as template and process and deleting empty pages as the end , before saving it as pdf'''
                # # OPTIMIZE: this part needs to be reviewd and optimised.
                #
                # # self.remove_empty_pages(self.doc)
                # self.remove_first_n_pages(self.doc, 4)

                # self.doc.save(self.existing_file_path)
                # new_path = self.add_space_in_footer(self.existing_file_path)
                # self.doc.save(new_path)
                # convert(new_path)


                # logger.info(f"full dict before fail - {dict}")
                # logger.info(f"the existing_file_path - {self.existing_file_path}")
                self.doc.save(self.existing_file_path)
                # logger.info(f"pdf file is - {self.pdf_path}")
                # logger.info("after the saving line...")
                # convert(self.pdf_path)
                convert(self.existing_file_path)
                # logger.info("after the convert line...")

                #self.doc.close()
            logger.info(
                f"END of the CYCLE.. process meta data getting method in thread....----   {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            return dict


        except Exception as e:
            import traceback

            logger.error(f"Error in get_save_vrms_data_to_get_report: {e}")
            traceback.print_exc()

            logger.error(f"Error in get_save_vrms_data_to_get_report: {e}")
            # raise  # Add this line to re-raise the exception

    # def get_input_dict(self):
    #     input_dict = self.get_save_vrms_data_to_get_report()
    #     doc = Document(self.existing_file_path)
    #     self.report_generate_obj.summary_table_work(doc,input_dict)

    # def remove_empty_pages(self,doc):
    #     try:
    #         paragraphs_to_remove = []
    #
    #         for para in doc.paragraphs:
    #             if not para.text.strip():
    #                 paragraphs_to_remove.append(para)
    #
    #         for para in paragraphs_to_remove:
    #             para.clear()
    #     except Exception as q:
    #         logger.error(f"Error in remove_empty_pages: {q}")

    # def remove_first_n_pages(self,doc, n):
    #     paragraphs_to_remove = []
    #
    #     # Iterate through paragraphs and identify those in the first n pages
    #     for i, para in enumerate(doc.paragraphs):
    #         # Check if the paragraph is in the first n pages
    #         if i < n:
    #             paragraphs_to_remove.append(para)
    #
    #     # Remove paragraphs from the document
    #     for para in paragraphs_to_remove:
    #         para.clear()

    # def add_space_in_footer(self,doc_path):
    #     try:
    #         doc = Document(doc_path)
    #         for section in doc.sections:
    #             footer = section.footer
    #             for paragraph in footer.paragraphs:
    #                 # Create a flag to identify when the "of" run is found
    #                 found_of = False
    #                 for run in paragraph.runs:
    #                     # If the "of" text is found and the flag is not already set
    #                     if 'of' in run.text and not found_of:
    #                         # Add a space after "of" if it doesn't already have one
    #                         if not run.text.endswith(' '):
    #                             run.text = run.text.replace('of', 'of ')
    #                         found_of = True  # Set the flag to True after finding "of"
    #                     # If the flag is set and the current run is a field (PAGE or NUMPAGES)
    #                     if found_of and ('PAGE' in run.text or 'NUMPAGES' in run.text):
    #                         # Add a space before the field if it doesn't already have one
    #                         if not run.text.startswith(' '):
    #                             run.text = ' ' + run.text
    #                         found_of = False  # Reset the flag after processing the field
    #
    #         # Save the modified document
    #         modified_doc_path = doc_path.replace('.docx', '_final.docx')
    #         doc.save(modified_doc_path)
    #         self.doc.close()
    #         return modified_doc_path
    #     except Exception as c:
    #         logger.error(f"error in footer empty line remove : line no 553-->{c}")

    @staticmethod
    def make_sure_path_exists(path):
        try:
            os.makedirs(path, exist_ok=True)
            # os.makedirs(path)
        except OSError as exception:
            if exception.errno != errno.EEXIST:
                raise


    def generate_axial_report_details_for_sensors(self):
        self.get_save_vrms_data_to_get_report()
        # for all sensors report with plots, remark, obs
        pass

    def get_euipments_details_for_summary(self):
        try:
            ko = self.metadata_extractor.get_hierarchy_level_with_equipment_details_from_metadata()
            # print(len(ko))
            result_list_str = []
            for item in ko:
                # result_string = f"{item['EquipmentName']}_{item['SensorMacID']}"
                # f"{item['EquipmentID']}_{item['EquipmentName']}_{item['EquipmentNumber']}_{item['SensorMacID']}"
                result_string = f"{item['EquipmentName']}"

                result_list_str.append(result_string)

            # print(len(result_list_str))
            # print(result_list_str)

            only_equipment_names = list(set(result_list_str))
            # print("*************************************************")
            # print("only_equipment_names",only_equipment_names)
            sorted_only_equipment_names = sorted(only_equipment_names)

            # print(sorted_only_equipment_names)
            # print("*************************************************")

            self.list_of_equips = sorted_only_equipment_names
            logger.info(f"{len(sorted_only_equipment_names)}  -------------this is count of equipments")
            return sorted_only_equipment_names
        except Exception as d:
            logger.error(f"Error in get_euipments_details_for_summary: {d}")

    def get_euipments_details_with_mac_id_for_summary_status(self):
        try:
            data = self.metadata_extractor.get_hierarchy_level_with_equipment_details_from_metadata()
            result_list_str = []
            if data:
                for item in data :
                    if item['SensorMacID'] not in self.exclude_mac_ids:
                        result_string = f"{item['EquipmentName']}-{item['SensorMacID']}"

                        # f"{item['EquipmentID']}_{item['EquipmentName']}_{item['EquipmentNumber']}_{item['SensorMacID']}"
                        # result_string = f"{item['EquipmentName']}"
                        # print(result_string)
                        result_list_str.append(result_string)
                    else:
                        logger.info(f"sensors not considering - {item['SensorMacID']}")

            # logger.info(result_list_str)
            sorted_list_of_equips_with_mac_id = sorted(result_list_str)

            # print(sorted_list_of_equips_with_mac_id)

            self.list_of_equips_with_mac_id = sorted_list_of_equips_with_mac_id

            logger.info(f"{len(sorted_list_of_equips_with_mac_id)}-------------this is count of mac ids")
            # logger.info(sorted_list_of_equips_with_mac_id)
            return sorted_list_of_equips_with_mac_id
        except Exception as d:
            logger.error(f"Error in get_euipments_details_with_mac_id_for_summary_status: {d}")

# if __name__ == "__main__":
#
#     # trends_json = "C:\\Trends_services\\Trends.json"
#     trends_json = "C:\\Vegam\\Trends_alert\\Trends_1.json"
#     try:
#
#         def get_current_ist_date():
#             try:
#                 ist_timezone = pytz.timezone('Asia/Kolkata')
#                 utc_now = datetime.utcnow()
#                 ist_now = utc_now.replace(tzinfo=pytz.utc).astimezone(ist_timezone)
#                 formatted_ist_date = ist_now.strftime("%d_%m_%Y")
#                 print(type(formatted_ist_date))
#
#                 return formatted_ist_date
#             except Exception as d:
#                 print(d)
#
#
#         # Example usage
#         print(get_current_ist_date())
#
#         starting_date = get_current_ist_date()
#
#         '''Check if the file exists'''
#         if not os.path.isfile(trends_json):
#             raise FileNotFoundError(f"File '{trends_json}' does not exist.")
#
#         '''Check if the file is empty'''
#         if os.path.getsize(trends_json) == 0:
#             raise ValueError(f"File '{trends_json}' is empty.")
#
#         with open(trends_json, 'r') as json_file:
#             Trends_Config_json = json.load(json_file)
#         area_check_lists = Trends_Config_json["trends"]["area_check"]
#         print(area_check_lists)
#         # ReportFacadeObj = ReportGeneratorFacade(trends_json, area_check_lists)
#         # ReportFacadeObj.make_report()
#         for area in area_check_lists:
#             logger.info(f"Report work is started for :{area}")
#             ReportFacadeObj = ReportGeneratorFacade(trends_json, area,starting_date)
#             ReportFacadeObj.make_report()
#             time.sleep(15)
#
#
#     except FileNotFoundError as e:
#         logger.error(str(e))
#     except ValueError as e:
#         logger.error(str(e))
#     except Exception as e:
#         logger.error(f"An unexpected error occurred: {str(e)}")

# if __name__ == "__main__":
#
#     # trends_json = "C:\\Trends_services\\Trends.json"
#     trends_json = "C:\\Vegam\\Trends_alert\\Trends_1.json"
#     try:
#         '''Check if the file exists'''
#         if not os.path.isfile(trends_json):
#             raise FileNotFoundError(f"File '{trends_json}' does not exist.")
#
#         '''Check if the file is empty'''
#         if os.path.getsize(trends_json) == 0:
#             raise ValueError(f"File '{trends_json}' is empty.")
#
#         ReportFacadeObj = ReportGeneratorFacade(trends_json)
#         ReportFacadeObj.make_report()
#
#         # j = ReportFacadeObj.get_save_vrms_data_trend_process()
#
#         # ReportFacadeObj.run() # call the method --> rename it
#
#     except FileNotFoundError as e:
#         logger.error(str(e))
#     except ValueError as e:
#         logger.error(str(e))
#     except Exception as e:
#         logger.error(f"An unexpected error occurred: {str(e)}")