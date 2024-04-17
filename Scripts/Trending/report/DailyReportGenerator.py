# 6/10/23 12:49 PM 00000
# 22/11/23 10:51 AM 00000
# import io
#
# import docx
# from PIL import Image
# import os
import logging
from docx2pdf import convert

# from docx.enum.text import WD_TAB_ALIGNMENT

# import re
import fitz
import os
import shutil


# from docx.oxml import parse_xml
# from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx import Document
from docx.shared import Pt
# from docx.shared import Pt
import pandas as pd
from docx.shared import Inches
from docx.oxml import OxmlElement, ns
# from docx.oxml.ns import qn
# import json
import numpy as np
from numpy import nan
# from scipy.stats import linregress
# import matplotlib.pyplot as plt
import matplotlib
from docx.shared import RGBColor

matplotlib.use('TkAgg')  # Use the TkAgg backend (or another suitable one)
# import time
# from datetime import datetime, timedelta
# from docx.enum.style import WD_STYLE_TYPE
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime, timedelta
# from Scripts.Trending.report.logger_file import logger

class DailyReportGenerator:
    def __init__(self, exist_file, pdf_path, original_template, vegam_logo, customer_name,
                 list_of_equips,list_of_equips_with_mac_id, outputs_folder_path, area_check,critical,warning,report_start_hour,report_start_minute,report_start_second,report_start_microsecond
                 ,report_end_hour,report_end_minute,report_end_second,report_end_microsecond,report_days_offset,percentage_threshold ):
        try:
            self.existing_file_path = exist_file  # "D://vegam_projects//GIT_works//Vegam_Analytics//Scripts//Trending//report//outputs//report_making.docx"
            self.original_file_path = original_template  # "D://vegam_projects//GIT_works//Vegam_Analytics//Scripts//Trending//report//inputs//template//template.docx"
            self.title_image_path = vegam_logo  # "D://vegam_projects//GIT_works//Vegam_Analytics//Scripts//Trending//report//inputs//images//vegam_logo.jpg"
            # self.metadata_extractor = ExtractMetaData(vmaint_api)
            self.pdf_path = pdf_path
            self.customer_name = customer_name
            #self.trends_json_path = trends_json_path
            # self.alert_json_path = alert_json_path
            # self.trends_date = trends_date
            self.list_mac_id_equips_weekly_report = list_of_equips
            self.list_of_equips_with_mac_id = list_of_equips_with_mac_id

            self.path_of_outputs_folder = outputs_folder_path
            # self.machine_pic_path = machine_pic_path
            # self.moving_avg_plot_path = moving_avg_plot_path
            # self.machine_imange_unavailable_pic_path = machine_imange_unavailable_pic_path
            self.equipment_list_with_sensor_id = []
            self.area_check = area_check
            self.overall_section_number = 1  # Initialize overall section number
            self.critical = critical
            self.warning = warning
            self.report_start_hour = report_start_hour
            self.report_start_minute = report_start_minute
            self.report_start_second = report_start_second
            self.report_start_microsecond = report_start_microsecond
            self.report_end_hour = report_end_hour
            self.report_end_minute = report_end_minute
            self.report_end_second = report_end_second
            self.report_end_microsecond = report_end_microsecond
            self.report_days_offset = report_days_offset
            self.percentage_threshold = percentage_threshold

        except Exception as d:
            logging.error(f"Error in line no 33: {d}")

    def duplicate_and_add_title(self, original_path, duplicate_path, title_image_path):
        try:
            shutil.copyfile(original_path, duplicate_path)
            doc = Document(duplicate_path)
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run()
            title_run.add_picture(title_image_path, width=Inches(3))
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.save(duplicate_path)
        except Exception as d:
            logging.error(f"Error in line no 46: {d}")

    def add_title_to_existing_doc(self, existing_path, title):
        try:
            doc = Document(existing_path)
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(title)
            title_run.font.name = "Calibri Bold"  #
            title_run.bold = True
            title_run.font.size = Pt(26)  # Pt(26)
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_page_break()
            doc.save(existing_path)

        except Exception as d:
            logging.error(f"Error in line no 66: {d}")

    def format_date(self, date_str):
        try:
            print(date_str, "00000")
            input_date = datetime.strptime(date_str, "%d/%m/%Y %H:%M:%S")
            # input_str, "%d/%m/%Y %H:%M:%S"
            output_str = input_date.strftime("%d %B %Y")
            # date_obj = datetime.strptime(date_str, "%d %B %Y")
            return output_str  # .strftime("%d %b %Y")
        except Exception as d:
            logging.error(f"Error in line no 77: {d}")

    def create_report_common(self,input_dict_list):
        try:
            # common_import = {"site_id": "001", "start_date": "01 November 2023", "end_date": "14 November 2023"} # need to change
            # site_id = common_import["site_id"]
            # result_get = self.meta_data_json

            if not os.path.exists(self.path_of_outputs_folder):
                try:
                    os.makedirs(self.path_of_outputs_folder)
                    logging.info(f'Created "{self.path_of_outputs_folder}" folder.')
                except Exception as e:
                    logging.error(f'Error creating "{self.path_of_outputs_folder}" folder: {str(e)}')

            if self.customer_name is not None:
                customer_name = self.customer_name
                end_time_api_time = datetime.now()

                # Date and time of the previous day
                start_time_api = datetime.now() - timedelta(days=1)
                start_time_api_time = start_time_api
                # Format dates in the desired format
                formatted_start_time = start_time_api_time.strftime("%d %B %Y")
                formatted_end_time = end_time_api_time.strftime("%d %B %Y")

                duration = f"{formatted_start_time} – {formatted_end_time}"

                title_text = f'\n\n\n\n\n{customer_name} \nCondition Monitoring Report - {self.area_check}\n{duration}'
                self.duplicate_and_add_title(self.original_file_path, self.existing_file_path, self.title_image_path)
                self.add_title_to_existing_doc(self.existing_file_path, title_text)


                doc = Document(self.existing_file_path)
                for paragraph in doc.paragraphs:
                    paragraph.alignment = 1  # 1 corresponds to center alignment

                #title_paragraph = doc.add_paragraph(title_text)

                # Set the alignment for the title paragraph to center
                #title_paragraph.alignment = 1  # 1 corresponds to center alignment

                #input_dict = {'MAC_ID': 'C258249CFE98', 'Velocity RMS - Radial Horizontal': 1.045885202390778, 'Velocity RMS - Radial Vertical': 1.9041070622570058, 'Velocity RMS - Axial': 2.376408072229036, 'Equipment ID': '1200MW_CT_U4_Fan9'}
                # input_dict_list = [{'MAC_ID': 'C258249CFE98', 'Velocity RMS - Radial Horizontal': 1.045885202390778,
                #   'Velocity RMS - Radial Vertical': 1.9041070622570058, 'Velocity RMS - Axial': 2.376408072229036,
                #   'Equipment ID': '1200MW_CT_U1_Fan1'},
                #  {'MAC_ID': 'D7E0DB929BF9', 'Velocity RMS - Radial Horizontal': 2.65018826385067,
                #   'Velocity RMS - Radial Vertical': 6.215167415258651, 'Velocity RMS - Axial': 4.597339692331943,
                #   'Equipment ID': '1200MW_CT_U4_Fan9'}]

                # # hack: commenting out next line for now as per current release needs.
                # #  [ only equipement details sensor wise , so first common pages are not required]'''
                # self.add_summary(doc) # adding blue box --> summary
                #self.equipment_overview_status_table_with_key(doc)  # equipments status for given time
                self.summary_table_work(doc,input_dict_list)
                # equipments status , date wise for given time--> making table # sp will start with api check [ 18th jan, 2024]
                # # self.add_heading_for_details_report(doc) #  Equipment Details:  adding this heading -> causing sensor data misplacement , not coming in one page.
                # doc.add_page_break()

                doc.save(self.existing_file_path)
                print("ok done..")






            else:
                logging.error("error in report making.")
        except Exception as d:
            logging.error(f"Error in create_report line no 110: {d}")




    def summary_table_work(self, doc,input_dict_list):
        try:
            '''doc = Document(self.existing_file_path)'''
            end_time_api_time = datetime.now().replace(hour=8, minute=0, second=0, microsecond=0)
            start_time_api = datetime.now() - timedelta(days=self.report_days_offset)
            start_time_api_time = start_time_api.replace(hour=8, minute=0, second=0, microsecond=0)

            # start_time = int(start_time_api_time.timestamp() * 1000)
            # end_time = int(end_time_api_time.timestamp() * 1000)
            formatted_start_time = start_time_api_time.strftime("%d-%m-%Y %I:%M %p")
            formatted_end_time = end_time_api_time.strftime("%d-%m-%Y %I:%M %p")

            header_text = f"\n Daily Report: {formatted_start_time} to {formatted_end_time}"
            header_paragraph = doc.add_paragraph(header_text)

            header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER #skb
            header_paragraph.alignment = WD_ALIGN_VERTICAL.CENTER

            run = header_paragraph.runs[0]
            run.font.size = Pt(18)
            run.font.bold = True
            combined_equipment_sensor_ids = self.list_mac_id_equips_weekly_report if len(
                self.list_mac_id_equips_weekly_report) != 0 else "no data"
            #print(combined_equipment_sensor_ids,']]]]]]]]]]]]')

            df = pd.DataFrame({'EquipmentID_Sensor_ID': combined_equipment_sensor_ids})
            equip_details_df = df[['EquipmentID_Sensor_ID']] ##this will give only equipment ids
            #print(self.list_of_equips_with_mac_id)
            #print(equip_details_df,'equip_details_df')

            combined_equipment_sensor_id = self.list_of_equips_with_mac_id if len(
                self.list_mac_id_equips_weekly_report) != 0 else "no data"
            df1 = pd.DataFrame({'EquipmentID_mac_ID': combined_equipment_sensor_id})
            equip_details_df_1 = df1[['EquipmentID_mac_ID']] ##this will give  equipment ids and macids
            #print(equip_details_df_1,'equip_details_df_1')



            # num_days = len(days_list)

            # Add a table with enough rows and columns
            # table = doc.add_table(rows=len(equip_details_df) + 2, cols=num_days + 2)
            table = doc.add_table(rows=len(equip_details_df_1) + 2, cols=7)

            # Add a table with enough rows and columns
            # table = doc.add_table(rows=10, cols=7)
            table.style = 'Table Grid'

            # Merge cells for the first row to have 7 columns with a specific structure
            cell_10 = table.cell(0, 0)
            cell_10_text = cell_10.text
            cell_11 = table.cell(0, 1)
            cell_11_text = cell_11.text = 'Velocity RMS - Radial Horizontal'
            cell_11.merge(table.cell(0, 2))

            cell_13 = table.cell(0, 3)
            cell_13_text = cell_13.text = 'Velocity RMS - Radial Vertical'
            cell_13.merge(table.cell(0, 4))
            cell_15 = table.cell(0, 5)
            cell_15_text = cell_15.text = 'Velocity RMS - Axial'
            cell_15.merge(table.cell(0, 6))
            cell_20 = table.cell(1, 0)
            cell_20_text = cell_20.text = 'EquipmentName_MacID'

            # Divide the merged cells into 2 columns (2nd, 3rd, 4th columns) in the second row
            for i in range(1, 6, 2):
                cell_sub = table.cell(1, i)
                cell_sub.text = 'Average (mm/s)'

                cell_sub = table.cell(1, i + 1)
                cell_sub.text = '% Data Beyond Critical'

                # Do not merge cells for the second row

            # Populate other rows as needed

            # Style the font size, alignment, and make text bold for specified cells
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if cell.text in ['Velocity RMS - Radial Horizontal', 'Velocity RMS - Radial Vertical',
                                             'Velocity RMS - Axial', 'EquipmentName_MacID', 'Average (mm/s)',
                                             '% Data Beyond Critical']:
                                run.font.size = Pt(11)  # Adjust font size as needed
                                run.bold = True  # Make the text bold
                                paragraph.alignment = 1  # Center alignment

                                # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #skb



            table.alignment = 1

            table.cell(0, 0).width = Inches(1.5)
            self.list_of_equips_with_mac_id = equip_details_df_1['EquipmentID_mac_ID'].to_list()
            logging.info(f"{len(self.list_of_equips_with_mac_id)} length of list_of_equips_with_mac_id")

            # self.equipment_list_with_sensor_id = equip_details_df['EquipmentID_Sensor_ID'].tolist()
            # print(len(self.equipment_list_with_sensor_id),'len(self.equipment_list_with_sensor_id)')

            if input_dict_list is not None:

                for input_dict in input_dict_list:
                    if input_dict is not None:
                        # print("input_dict.get('EquipmentName_MacID')",input_dict.get('EquipmentName_MacID'))
                        for equip, equipment_id in enumerate(self.list_of_equips_with_mac_id):
                            table.cell(equip + 2, 0).text = f"{equipment_id}"
                            #print(equipment_id,'////////////////')
                            for equip_row in range(2, len(equip_details_df_1) + 2):
                                # print("input_dict.get('EquipmentName_MacID')",input_dict.get('EquipmentName_MacID'))
                                # print("table.cell(equip_row, 0).text",table.cell(equip_row, 0).text)
                                #print(table.cell(equip_row,0).text,'23456789')##Common_TT_BC1_Conveyor_D2930832E760
                                if table.cell(equip_row, 0).text == input_dict.get('EquipmentName_MacID', ''):

                                    cell = table.cell(equip_row, 0)
                                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER  # skb

                                    # Enable text wrapping for the first run in the first paragraph of the cell
                                    cell.paragraphs[0].runs[0].font.wrap_text = True

                                    for col_idx, column_name in enumerate(
                                            ['Velocity RMS - Radial Horizontal', 'Velocity RMS - Radial Vertical',
                                             'Velocity RMS - Axial'], start=1):
                                        value = input_dict.get(column_name, '')
                                        if isinstance(value, np.float64):
                                            value = float(value)
                                        sub_col_idx = col_idx * 2 - 1

                                        table.cell(equip_row, sub_col_idx).text = str(value)
                                        table.cell(equip_row, sub_col_idx).paragraphs[
                                            0].alignment = 1  # Center alignment

                                        # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #skb

                                        # Set the '% Data Beyond Critical' values in the next sub-column
                                        percent_beyond_critical_col = sub_col_idx + 1
                                        critical_column_name = f'{column_name}_% Data Beyond Critical'
                                        percent_beyond_critical_value = input_dict.get(critical_column_name, '')
                                        #print(percent_beyond_critical_value,'------------')

                                        #print(type(percent_beyond_critical_value),'\\\\\\\\\\\\')
                                        if isinstance(percent_beyond_critical_value, np.float64):
                                            percent_beyond_critical_value = float(percent_beyond_critical_value)

                                        table.cell(equip_row, percent_beyond_critical_col).text = str(
                                            percent_beyond_critical_value)
                                        table.cell(equip_row, percent_beyond_critical_col).paragraphs[
                                            0].alignment = 1  # Center alignment

                                        # cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER #skb


                                        # print("before if type percent_beyond_critical_value - ",
                                        #       type(percent_beyond_critical_value))
                                        if percent_beyond_critical_value not in ["OFF", "-", "nan", "NaN", None, np.NaN, nan,
                                                                                               '<NA>'] and percent_beyond_critical_value > self.percentage_threshold:

                                            cell = table.cell(equip_row, percent_beyond_critical_col)  # Access the correct cell
                                            run = cell.paragraphs[0].runs[0]
                                            # Set the font color to red
                                            font = run.font
                                            font.color.rgb = RGBColor(255, 0, 0)
                                            #print('red coloring done')
                                        # print("after if type percent_beyond_critical_value - ",type(percent_beyond_critical_value)) #skb

                                        # print("before if type value - ", type(value))  # skb
                                        if isinstance(value, (int,float)) and value not in ["OFF", "-", "nan", "NaN", None, np.NaN, nan,
                                                                                               '<NA>'] and value is not nan and value > self.critical:
                                            cell = table.cell(equip_row,
                                                              sub_col_idx)  # Access the correct cell
                                            run = cell.paragraphs[0].runs[0]
                                            # Set the font color to red
                                            font = run.font
                                            font.color.rgb = RGBColor(255, 0, 0)
                                        # print("after if type value - ", type(value))

                    # Leave the last row empty
                    table.rows[len(equip_details_df_1) + 1].height = Pt(0)
            else:
                logging.error("Input dictionary list is None.")


            footer_text = "\n\nNOTE:    No Data Available : ' - '\n"\
                          "               Machine OFF : ' OFF '"
            footer_paragraph = doc.add_paragraph(footer_text)

            footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # skb
            footer_paragraph.paragraph_format.space_after = Pt(0)
            footer_paragraph.paragraph_format.space_before = Inches(0.25)
            # footer_paragraph.alignment = WD_ALIGN_VERTICAL.CENTER

            run = footer_paragraph.runs[0]
            run.font.size = Pt(11)
            run.font.bold = True
            # footer_paragraph.paragraph_format.line_spacing = Pt(12)
            highlight_text = "NOTE:"

            for paragraph in doc.paragraphs:
                if highlight_text in paragraph.text:
                    # Split the paragraph text by the highlight_text
                    parts = paragraph.text.split(highlight_text)

                    # Create a new run for the first part (before the highlight_text)
                    run = paragraph.runs[0]
                    run.text = parts[0]

                    # Add a new run for the highlight_text and set its color to red
                    paragraph.add_run(highlight_text).font.color.rgb = RGBColor(255, 0, 0)

                    # Create a new run for the rest of the paragraph (after the highlight_text)
                    if len(parts) > 1:
                        run = paragraph.add_run(parts[1])
                        run.font.color.rgb = RGBColor(0, 0, 0)

        except Exception as e:
            import traceback
            traceback.print_exc()
            logging.error(f"An error occurred: {e}")

    # content works.header footer...

    def create_element(self, name):
        return OxmlElement(name)

    def create_attribute(self, element, name, value):
        try:
            element.set(ns.qn(name), value)
        except Exception as e:
            # Handle the exception as per your requirements
            logging.error(f"An error occurred while setting attribute: {e}")


    def add_page_number(self, paragraph):
        try:

            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Additional text
            additional_text_run = paragraph.add_run()
            additional_text = self.create_element('w:t')
            self.create_attribute(additional_text, 'xml:space', 'preserve')
            additional_text.text = 'Vegam Solutions Pte Ltd         www.@vegam.co      Confidential    Not for redistribution                 Page '
            additional_text_run._r.append(additional_text)

            # Add the page number field
            fldChar1 = self.create_element('w:fldChar')
            self.create_attribute(fldChar1, 'w:fldCharType', 'begin')

            instrText = self.create_element('w:instrText')
            self.create_attribute(instrText, 'xml:space', 'preserve')
            instrText.text = "PAGE"

            fldChar2 = self.create_element('w:fldChar')
            self.create_attribute(fldChar2, 'w:fldCharType', 'end')

            additional_text_run._r.append(fldChar1)
            additional_text_run._r.append(instrText)
            additional_text_run._r.append(fldChar2)

            # Add space between additional text and "Page" field
            space_run1 = paragraph.add_run()
            space_text1 = self.create_element('w:t')
            self.create_attribute(space_text1, 'xml:space', 'preserve')
            space_text1.text = ' '  # Add a space
            space_run1._r.append(space_text1)

            # Add "of" between "Page" and total page count
            of_run = paragraph.add_run()
            of_text = self.create_element('w:t')
            self.create_attribute(of_text, 'xml:space', 'preserve')
            of_text.text = 'of'
            of_run._r.append(of_text)

            # Add space between "of" and total page count
            space_run2 = paragraph.add_run()
            space_text2 = self.create_element('w:t')
            self.create_attribute(space_text2, 'xml:space', 'preserve')
            space_text2.text = ' '  # Add a space
            space_run2._r.append(space_text2)

            # Add the total page count field
            fldChar3 = self.create_element('w:fldChar')
            self.create_attribute(fldChar3, 'w:fldCharType', 'begin')

            instrText2 = self.create_element('w:instrText')
            self.create_attribute(instrText2, 'xml:space', 'preserve')
            instrText2.text = "NUMPAGES"

            fldChar4 = self.create_element('w:fldChar')
            self.create_attribute(fldChar4, 'w:fldCharType', 'end')

            space_run2._r.append(fldChar3)
            space_run2._r.append(instrText2)
            space_run2._r.append(fldChar4)

            # Set font size and style for the entire line
            for run in [additional_text_run, space_run1, of_run, space_run2]:
                run.font.size = Pt(10)
                run.font.name = "Calibri"
                run.font.bold = True

        except Exception as d:
            logging.error(f"Error in add_page_number: {d}")

    # add page number to content page
    def extract_page_numbers(self, search_terms):
        try:
            doc = fitz.open(self.pdf_path)

            page_numbers = {}

            for term in search_terms:
                term = term.lower()
                for page_number in range(doc.page_count):
                    page = doc[page_number]
                    text = page.get_text("text")
                    if term in text.lower():
                        page_numbers[term] = page_number + 1
                        break

            doc.close()

        except Exception as e:
            logging.error(f"An error occurred while extracting page numbers: {e}")
            return {}

        return page_numbers

    def find_page_number(self, pdf_path, search_term):
        try:
            doc = fitz.open(pdf_path)
            #print(search_term,'im search termmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmmm')


            for page_number in range(1, doc.page_count + 1):  # Start from page 1 (exclude CONTENTS page)
                page = doc[page_number - 1]
                text = page.get_text("text")

                # Check if the search term is present in the text
                if search_term.lower() in text.lower():
                    logging.critical(f"Search term '{search_term}' found on page {page_number}")
                    doc.close()
                    return page_number

            logging.critical(f"Search term '{search_term}' not found in the document.")
            doc.close()
            return None  # Return None if the search term is not found in the document

        except Exception as e:
            logging.error(f"Error in find_page_number: {e}")

    def add_page_numbers_to_contents(self, doc, pdf_path):
        try:
            contents_section = None

            # Iterate through paragraphs to find the CONTENTS section
            for paragraph in doc.paragraphs:
                if 'CONTENTS' in paragraph.text:
                    contents_section = paragraph
                    break


            if contents_section is not None:
                # Iterate through lines in the CONTENTS section
                for run in contents_section.runs:
                    # Add a colon to each line
                    line_with_colon = f'{run.text}:'
                    print(line_with_colon,'..................................................................')

                    # Find the page number for each line
                    page_number = self.find_page_number(pdf_path, line_with_colon)

                    if page_number is not None:
                        # Append the page number directly to the line along with dots
                        dots = '.' * (100 - len(run.text) - len(str(page_number)))
                        run.text = f'{run.text} {dots} {page_number}'
        except Exception as e:
            logging.error(f"Error in add_page_numbers_to_contents: {e}")

    def content_page_works(self, doc):
        try:
            # footer...
            # Iterate through each section in the document
            for section in doc.sections:
                # Check if the section has a footer
                if section.footer:
                    footer = section.footer
                    # Check if the footer has paragraphs and the page number is not already present
                    if footer.paragraphs and "Page " not in footer.paragraphs[0].text:
                        paragraph = footer.paragraphs[0]
                        # Remove leading spaces from each run in the paragraph
                        for run in paragraph.runs:
                            run.text = run.text.lstrip()
                        # Set the alignment for the paragraph
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        # Add page number to the paragraph
                        self.add_page_number(paragraph)
                    elif not footer.paragraphs:
                        # If the footer does not have paragraphs, create one and add the page number
                        paragraph = footer.add_paragraph()
                        for run in paragraph.runs:
                            run.text = run.text.lstrip()

                        self.add_page_number(paragraph)
                    # No else needed, if "Page " is already in the footer, nothing to do

            # Save the document with any changes made
            doc.save(self.existing_file_path)

            # Save the document with any changes made
            # doc.save(self.existing_file_path)
            convert(self.existing_file_path, self.pdf_path)
            self.add_page_numbers_to_contents(doc, self.pdf_path)
            doc.save(self.existing_file_path)



        except Exception as e:
            logging.error(f"Error in content_page_works: {e}")
